"""Content processing pipeline.

Converts uploaded or captured content into searchable chunks:
- PDF: extract text and embedded images.
- Image: keep the original image chunk and optionally add OCR text.
- Audio: transcribe via the configured provider and create timestamped chunks.
- Video: transcribe via the configured provider and optionally capture a frame.
- Office: extract text plus basic document structure metadata.
- Web: extract article text and optionally capture a page screenshot.
"""

import asyncio
import base64
import hashlib
import logging
import mimetypes
import traceback
import uuid
from datetime import datetime, timezone
from pathlib import Path

import numpy as np
import httpx
from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.logging import get_logger
from app.models.models import Content, ContentChunk, FunctionBindingConfig, ProviderConfig
from app.core.crypto import crypto_service

settings = get_settings()
logger = get_logger(__name__)


# ── PDF 文字提取 ──

def _extract_pdf_text(path: Path) -> list[tuple[int, str]]:
    """提取 PDF 文字，返回 (page_number, text) 列表"""
    import fitz

    doc = fitz.open(path)
    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            pages.append((page_num + 1, text))
    doc.close()
    return pages


# ── PDF 图片提取 ──

def _extract_pdf_images(path: Path, output_dir: Path) -> list[tuple[int, str]]:
    """提取 PDF 内嵌图片，返回 (page_number, image_path) 列表"""
    import fitz

    doc = fitz.open(path)
    output_dir.mkdir(parents=True, exist_ok=True)
    images = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        for img_index, img in enumerate(page.get_images()):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n > 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                filename = f"page{page_num + 1}_img{img_index}.png"
                img_path = output_dir / filename
                pix.save(str(img_path))
                images.append((page_num + 1, filename))
                pix = None
            except Exception:
                continue

    doc.close()
    return images


# ── Office 文档解析 ──

def _extract_docx(path: Path) -> tuple[str, dict]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    headings: list[dict] = []
    paragraph_count = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_count += 1
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            level_text = style_name.split()[-1]
            level = int(level_text) if level_text.isdigit() else None
            headings.append({"text": text, "level": level, "paragraph_index": index})
        parts.append(text)

    tables: list[dict] = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
                parts.append(" | ".join(cell for cell in cells if cell))
        tables.append({
            "table_index": table_index,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "preview": rows[:5],
        })

    metadata = {
        "format": "docx",
        "paragraph_count": paragraph_count,
        "headings": headings,
        "tables": tables,
    }
    return "\n".join(parts), metadata


def _extract_xlsx(path: Path) -> tuple[str, dict]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True)
    parts = []
    sheets: list[dict] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        non_empty_rows = 0
        preview = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                non_empty_rows += 1
                if len(preview) < 5:
                    preview.append(cells)
                parts.append(" | ".join(cells))
        sheets.append({
            "name": sheet_name,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "non_empty_rows": non_empty_rows,
            "preview": preview,
        })
    wb.close()
    return "\n".join(parts), {"format": "xlsx", "sheets": sheets}


# -- Image OCR --

async def _get_function_binding(db: AsyncSession, function: str) -> tuple[ProviderConfig, str] | None:
    binding_result = await db.execute(
        select(FunctionBindingConfig).where(FunctionBindingConfig.function == function)
    )
    binding = binding_result.scalar_one_or_none()
    if binding and binding.provider_id and binding.model:
        provider_result = await db.execute(
            select(ProviderConfig).where(
                ProviderConfig.id == binding.provider_id,
                ProviderConfig.is_active == True,
            )
        )
        provider = provider_result.scalar_one_or_none()
        if provider is not None:
            return provider, binding.model

    result = await db.execute(
        select(ProviderConfig).where(ProviderConfig.is_active == True)
    )
    for provider in result.scalars().all():
        models = provider.default_models or {}
        model = models.get(function)
        if model:
            return provider, model

    return None


async def _get_ocr_binding(db: AsyncSession) -> tuple[ProviderConfig, str] | None:
    return await _get_function_binding(db, "ocr")


async def _ocr_image(path: Path, db: AsyncSession | None = None) -> str:
    if db is None:
        return ""

    binding = await _get_ocr_binding(db)
    if binding is None:
        logger.info("OCR provider is not configured; skipping image text extraction")
        return ""

    if not path.exists():
        logger.warning(f"OCR image file does not exist: {path}")
        return ""

    provider, model = binding
    api_key = ""
    if provider.api_key_encrypted:
        try:
            api_key = crypto_service.decrypt(provider.api_key_encrypted)
        except Exception:
            logger.warning(f"OCR API key decrypt failed - provider_id={provider.id}", exc_info=True)
            return ""

    try:
        from openai import AsyncOpenAI

        mime_type = mimetypes.guess_type(path.name)[0] or "image/png"
        image_data = base64.b64encode(path.read_bytes()).decode("ascii")
        data_url = f"data:{mime_type};base64,{image_data}"
        base_url = provider.base_url or settings.openai_base_url
        client = AsyncOpenAI(api_key=api_key or "no-key", base_url=base_url)
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Extract all readable text from this image. "
                                "Return only the extracted text. If there is no readable text, return an empty string."
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            temperature=0,
        )
        content = response.choices[0].message.content if response.choices else ""
        if isinstance(content, str):
            return content.strip()
        if isinstance(content, list):
            parts = [
                item.get("text", "")
                for item in content
                if isinstance(item, dict) and item.get("type") == "text"
            ]
            return "\n".join(part for part in parts if part).strip()
        return ""
    except Exception:
        logger.warning(
            f"OCR extraction failed - provider_id={provider.id}, model={model}, path={path}",
            exc_info=True,
        )
        return ""


# -- Audio transcription --

def _normalize_transcript_segments(response) -> list[dict]:
    segments = getattr(response, "segments", None)
    if segments:
        normalized = []
        for seg in segments:
            text = getattr(seg, "text", None)
            if text is None and isinstance(seg, dict):
                text = seg.get("text")
            text = (text or "").strip()
            if not text:
                continue
            start = getattr(seg, "start", None)
            end = getattr(seg, "end", None)
            if isinstance(seg, dict):
                start = seg.get("start", start)
                end = seg.get("end", end)
            normalized.append({"text": text, "start": start, "end": end})
        if normalized:
            return normalized

    text = getattr(response, "text", None)
    if text is None and isinstance(response, dict):
        text = response.get("text")
    text = (text or "").strip()
    if not text:
        return []
    return [{"text": text, "start": None, "end": None}]


async def _transcribe_with_faster_whisper(path: Path, model: str, extra_params: dict | None = None) -> list[dict]:
    """Run local faster-whisper transcription in a worker thread."""
    extra_params = extra_params or {}

    def run() -> list[dict]:
        from faster_whisper import WhisperModel

        model_size_or_path = extra_params.get("model_path") or model
        device = extra_params.get("device", "auto")
        compute_type = extra_params.get("compute_type", "default")
        language = extra_params.get("language") or None
        beam_size = int(extra_params.get("beam_size", 5))

        whisper_model = WhisperModel(
            model_size_or_path,
            device=device,
            compute_type=compute_type,
        )
        segments, _info = whisper_model.transcribe(
            str(path),
            language=language,
            beam_size=beam_size,
        )
        return [
            {"text": segment.text.strip(), "start": segment.start, "end": segment.end}
            for segment in segments
            if segment.text and segment.text.strip()
        ]

    return await asyncio.to_thread(run)


async def _transcribe_audio(path: Path, db: AsyncSession | None = None) -> list[dict]:
    """Return transcript segments as [{"text": "...", "start": 0.0, "end": 2.5}, ...]."""
    if db is None:
        return []

    binding = await _get_function_binding(db, "transcribe")
    if binding is None:
        logger.info("Transcription provider is not configured; skipping audio text extraction")
        return []

    if not path.exists():
        logger.warning(f"Audio file does not exist: {path}")
        return []

    provider, model = binding
    extra_params = provider.extra_params or {}
    binding_backend = extra_params.get("transcribe_backend")
    if binding_backend == "faster_whisper" or provider.base_url == "local:faster-whisper":
        try:
            return await _transcribe_with_faster_whisper(path, model, extra_params)
        except Exception:
            logger.warning(
                f"Local faster-whisper transcription failed - provider_id={provider.id}, model={model}, path={path}",
                exc_info=True,
            )
            return []

    api_key = ""
    if provider.api_key_encrypted:
        try:
            api_key = crypto_service.decrypt(provider.api_key_encrypted)
        except Exception:
            logger.warning(f"Transcription API key decrypt failed - provider_id={provider.id}", exc_info=True)
            return []

    try:
        from openai import AsyncOpenAI

        base_url = provider.base_url or settings.openai_base_url
        client = AsyncOpenAI(api_key=api_key or "no-key", base_url=base_url)
        with path.open("rb") as audio_file:
            try:
                response = await client.audio.transcriptions.create(
                    model=model,
                    file=audio_file,
                    response_format="verbose_json",
                    timestamp_granularities=["segment"],
                )
            except TypeError:
                audio_file.seek(0)
                response = await client.audio.transcriptions.create(
                    model=model,
                    file=audio_file,
                    response_format="json",
                )
        return _normalize_transcript_segments(response)
    except Exception:
        logger.warning(
            f"Audio transcription failed - provider_id={provider.id}, model={model}, path={path}",
            exc_info=True,
        )
        return []


# -- Video processing --

async def _extract_video_screenshots(path: Path, output_dir: Path) -> list[str]:
    import shutil

    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        logger.info("ffmpeg is not available; skipping video screenshots")
        return []

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "frame_0001.jpg"
    try:
        proc = await asyncio.create_subprocess_exec(
            ffmpeg,
            "-y",
            "-ss",
            "1",
            "-i",
            str(path),
            "-frames:v",
            "1",
            str(output_path),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=20)
        if proc.returncode != 0 or not output_path.exists():
            logger.warning(
                f"Video screenshot extraction failed - path={path}, "
                f"stdout={stdout.decode(errors='ignore')[:300]}, "
                f"stderr={stderr.decode(errors='ignore')[:300]}"
            )
            return []

        storage_root = Path(settings.file_storage_root).resolve()
        try:
            return [output_path.resolve().relative_to(storage_root).as_posix()]
        except ValueError:
            return [str(output_path)]
    except Exception:
        logger.warning(f"Video screenshot extraction failed - path={path}", exc_info=True)
        return []


async def _process_video(
    path: Path,
    db: AsyncSession | None = None,
    screenshot_dir: Path | None = None,
) -> tuple[list[dict], list[str]]:
    """Return transcript segments and optional screenshot paths."""
    segments = await _transcribe_audio(path, db)
    screenshots = await _extract_video_screenshots(path, screenshot_dir) if screenshot_dir else []
    return (segments, screenshots)


# ── 网页抓取 ──

async def _extract_web(url: str) -> str:
    import trafilatura

    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        return ""
    result = trafilatura.extract(downloaded)
    return result or ""


async def _capture_web_screenshot(url: str, output_dir: Path) -> str | None:
    try:
        from playwright.async_api import async_playwright
    except Exception:
        logger.info("Playwright is not installed; skipping web screenshot capture")
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "web_capture.png"
    try:
        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            page = await browser.new_page(viewport={"width": 1365, "height": 900})
            await page.goto(url, wait_until="networkidle", timeout=15000)
            await page.screenshot(path=str(output_path), full_page=True)
            await browser.close()

        storage_root = Path(settings.file_storage_root).resolve()
        try:
            return output_path.resolve().relative_to(storage_root).as_posix()
        except ValueError:
            return str(output_path)
    except Exception:
        logger.warning(f"Web screenshot capture failed - url={url}", exc_info=True)
        return None


# ── 主处理类 ──

class ContentProcessService:
    """内容处理管道：解析 → 分块 → 向量化"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def process(self, content_id: str | None = None, content: Content | None = None, keep_embedded: bool = True) -> Content:
        """处理单个 Content：解析文本 → 语义分块 → 生成嵌入（完整流程）
        
        :param keep_embedded: 是否保留已经有 embedding 的 chunks，默认为 True
        """
        import time
        start_time = time.time()
        
        if content is None:
            result = await self.db.execute(select(Content).where(Content.id == content_id))
            content = result.scalar_one_or_none()
            if content is None:
                raise ValueError(f"Content {content_id} not found")

        logger.info(f"开始完整处理内容 - content_id={content.id}, type={content.content_type}, title={content.title}")
        content.processing_status = "processing"
        await self.db.flush()

        try:
            await self._dispatch_and_chunk(content, keep_embedded=keep_embedded)
            content.processing_status = "completed"
            await self.db.commit()
            elapsed_time = (time.time() - start_time) * 1000
            logger.info(f"内容完整处理完成 - content_id={content.id}, 耗时: {elapsed_time:.2f}ms")
        except Exception:
            content.processing_status = "failed"
            content.processing_error = traceback.format_exc()
            logger.error(f"内容处理失败 - content_id={content.id}: {content.processing_error}")
            await self.db.commit()
            raise

        await self.db.refresh(content)
        return content

    async def chunk(self, content_id: str | None = None, content: Content | None = None) -> Content:
        """智能分块：解析文本 → 语义分块（不生成嵌入）
        
        :param content_id: 内容ID
        :param content: 内容对象（如果已获取）
        :return: 更新后的 Content 对象
        """
        import time
        start_time = time.time()
        
        if content is None:
            result = await self.db.execute(select(Content).where(Content.id == content_id))
            content = result.scalar_one_or_none()
            if content is None:
                raise ValueError(f"Content {content_id} not found")

        logger.info(f"开始智能分块 - content_id={content.id}, type={content.content_type}, title={content.title}")
        content.processing_status = "chunking"
        await self.db.flush()

        try:
            await self._dispatch_and_chunk(content, keep_embedded=True)
            content.processing_status = "chunked"
            await self.db.commit()
            elapsed_time = (time.time() - start_time) * 1000
            logger.info(f"智能分块完成 - content_id={content.id}, 耗时: {elapsed_time:.2f}ms")
        except Exception:
            content.processing_status = "failed"
            content.processing_error = traceback.format_exc()
            logger.error(f"分块失败 - content_id={content.id}: {content.processing_error}")
            await self.db.commit()
            raise

        await self.db.refresh(content)
        return content

    async def embed(self, content_id: str | None = None, content: Content | None = None) -> Content:
        """生成嵌入向量（仅对已有分块进行向量化）
        
        :param content_id: 内容ID
        :param content: 内容对象（如果已获取）
        :return: 更新后的 Content 对象
        """
        import time
        start_time = time.time()
        
        if content is None:
            result = await self.db.execute(select(Content).where(Content.id == content_id))
            content = result.scalar_one_or_none()
            if content is None:
                raise ValueError(f"Content {content_id} not found")

        logger.info(f"开始生成嵌入向量 - content_id={content.id}, title={content.title}")
        content.processing_status = "embedding"
        await self.db.flush()

        try:
            await self._embed_chunks_batched(content.id)
            elapsed_time = (time.time() - start_time) * 1000
            logger.info(f"嵌入向量生成完成 - content_id={content.id}, 耗时: {elapsed_time:.2f}ms")
        except Exception:
            content.processing_status = "failed"
            content.processing_error = traceback.format_exc()
            logger.error(f"嵌入失败 - content_id={content.id}: {content.processing_error}")
            await self.db.commit()
            raise

        await self.db.refresh(content)
        return content

    async def _dispatch_and_chunk(self, content: Content, keep_embedded: bool = True) -> None:
        """根据内容类型分派处理并执行分块"""
        content_id = str(content.id)
        
        await self._clear_old_chunks(content_id, keep_embedded)

        if content.content_type == "pdf":
            await self._process_pdf(content)
        elif content.content_type == "image":
            await self._process_image(content)
        elif content.content_type == "audio":
            await self._process_audio(content)
        elif content.content_type == "video":
            await self._process_video(content)
        elif content.content_type == "doc":
            await self._process_doc(content)
        elif content.content_type == "note":
            await self._process_note(content)
        elif content.content_type == "web":
            await self._process_web(content)
        else:
            logger.warning(f"未知内容类型 - content_id={content_id}, type={content.content_type}")

    async def _clear_old_chunks(self, content_id: str, keep_embedded: bool) -> None:
        """清除旧的分块"""
        if keep_embedded:
            # 只删除未嵌入的 chunks
            await self.db.execute(
                delete(ContentChunk).where(
                    ContentChunk.content_id == content_id,
                    ContentChunk.embedding.is_(None)
                )
            )
            logger.info(f"仅删除未嵌入的旧 chunks - content_id={content_id}")
        else:
            # 删除所有旧 chunks
            await self.db.execute(
                delete(ContentChunk).where(ContentChunk.content_id == content_id)
            )
            logger.info(f"删除所有旧 chunks - content_id={content_id}")
        await self.db.flush()

    async def _embed_chunks_batched(self, content_id, chunk_type: str = "text") -> None:
        """Generate embeddings for pending chunks; text chunks are batched.

        处理流程：
        1. 收集待处理 chunks（text 和 image 分别处理）
        2. 文本 chunks 按 32 条/批 批量调用嵌入 API
        3. 图片 chunks 单独调用嵌入 API
        4. 根据成功/失败数量更新 content 状态
        5. 计算内容级向量（分块向量的平均值）
        """
        logger.info(f"开始为内容生成嵌入向量 - content_id={content_id}")

        result = await self.db.execute(
            select(ContentChunk).where(
                ContentChunk.content_id == content_id,
                ContentChunk.embedding.is_(None),
            )
        )
        chunks = result.scalars().all()

        if not chunks:
            logger.info(f"没有需要处理的 chunks - content_id={content_id}")
            await self._compute_content_embedding(content_id)
            return

        text_chunks = [chunk for chunk in chunks if chunk.chunk_text and chunk.chunk_type != "image"]
        image_chunks = [chunk for chunk in chunks if chunk.chunk_type == "image" and chunk.image_path]
        skipped = len(chunks) - len(text_chunks) - len(image_chunks)

        logger.info(
            f"处理嵌入向量 - content_id={content_id}, "
            f"总 chunks: {len(chunks)}, 文本: {len(text_chunks)}, 图片: {len(image_chunks)}, 跳过: {skipped}"
        )

        from app.services.embedding import embed_image, embed_images, embed_texts

        success = 0
        failed = 0

        # 处理文本 chunks
        batch_size = 32
        total_text_batches = (len(text_chunks) + batch_size - 1) // batch_size if text_chunks else 0
        for batch_idx, start in enumerate(range(0, len(text_chunks), batch_size)):
            batch = text_chunks[start:start + batch_size]
            logger.info(
                f"处理文本批次 {batch_idx + 1}/{total_text_batches} - "
                f"content_id={content_id}, 数量: {len(batch)}"
            )
            try:
                vecs = await embed_texts(self.db, [chunk.chunk_text or "" for chunk in batch])
            except Exception as e:
                msg = f"embed_texts error for content {content_id}, batch_start={start}: {e}"
                logger.error(msg, exc_info=True)
                failed += len(batch)
                continue

            for chunk, vec in zip(batch, vecs):
                if vec:
                    chunk.embedding = vec
                    chunk.embedding_type = "text"
                    success += 1
                else:
                    msg = f"embed_texts returned None for chunk {chunk.id}"
                    logger.warning(msg)
                    failed += 1

            await self.db.commit()

        # 处理图片 chunks（批量处理）
        if image_chunks:
            image_batch_size = 32
            total_image_batches = (len(image_chunks) + image_batch_size - 1) // image_batch_size
            
            for batch_idx, start in enumerate(range(0, len(image_chunks), image_batch_size)):
                batch = image_chunks[start:start + image_batch_size]
                logger.info(
                    f"处理图片批次 {batch_idx + 1}/{total_image_batches} - "
                    f"content_id={content_id}, 数量={len(batch)}"
                )
                try:
                    vecs = await embed_images(self.db, [chunk.image_path for chunk in batch])
                except Exception as e:
                    msg = f"embed_images error for content {content_id}, batch_start={start}: {e}"
                    logger.error(msg, exc_info=True)
                    failed += len(batch)
                    continue
                
                for chunk, vec in zip(batch, vecs):
                    if vec:
                        chunk.embedding = vec
                        chunk.embedding_type = "image"
                        success += 1
                    else:
                        msg = f"embed_images returned None for chunk {chunk.id}"
                        logger.warning(msg)
                        failed += 1
                
                await self.db.commit()

        failed += skipped

        # 更新 Content 状态
        content_result = await self.db.execute(
            select(Content).where(Content.id == content_id)
        )
        content = content_result.scalar_one_or_none()
        if content:
            if failed == 0 and success > 0:
                content.processing_status = "completed"
                content.processed_at = datetime.now(timezone.utc)
                logger.info(f"内容嵌入处理成功 - content_id={content_id}")
            elif success == 0:
                content.processing_status = "failed"
                content.processing_error = f"所有 chunks 嵌入失败 ({failed}/{len(chunks)})"
                logger.error(f"内容嵌入处理失败 - content_id={content_id}, success={success}, failed={failed}")
            else:
                content.processing_status = "partial"
                content.processing_error = f"部分失败 ({failed}/{len(chunks)})"
                logger.warning(f"内容嵌入处理部分成功 - content_id={content_id}, success={success}, failed={failed}")

        # 计算内容级向量（分块向量的平均值）
        if success > 0:
            await self._compute_content_embedding(content_id)

        summary = (
            f"嵌入处理完成 - content_id={content_id}, "
            f"success={success}, failed={failed}, total={len(chunks)}, skipped={skipped}"
        )
        logger.info(summary)
        await self.db.commit()

    async def _compute_content_embedding(self, content_id: str) -> None:
        """计算内容级向量：对所有分块向量求平均"""
        logger.info(f"计算内容级向量 - content_id={content_id}")
        
        result = await self.db.execute(
            select(ContentChunk).where(
                ContentChunk.content_id == content_id,
                ContentChunk.embedding.is_not(None)
            )
        )
        chunks_with_embedding = result.scalars().all()
        
        if not chunks_with_embedding:
            logger.info(f"没有可用的分块向量用于计算内容级向量 - content_id={content_id}")
            return

        # 获取所有分块向量并求平均
        vectors = [np.array(chunk.embedding) for chunk in chunks_with_embedding]
        if vectors:
            avg_vector = np.mean(vectors, axis=0)
            
            # 获取嵌入类型（优先取文本类型，如果只有图片则取图片类型）
            embedding_type = None
            has_text = any(c.embedding_type == "text" for c in chunks_with_embedding)
            has_image = any(c.embedding_type == "image" for c in chunks_with_embedding)
            if has_text:
                embedding_type = "text"
            elif has_image:
                embedding_type = "image"

            # 更新内容级向量
            content_result = await self.db.execute(
                select(Content).where(Content.id == content_id)
            )
            content = content_result.scalar_one_or_none()
            if content:
                content.embedding = avg_vector.tolist()
                content.embedding_type = embedding_type
                logger.info(f"内容级向量计算完成 - content_id={content_id}, 基于 {len(vectors)} 个分块")

    def _resolve_path(self, file_path: str | None) -> str:
        if not file_path:
            raise ValueError("file_path is empty")
        root = Path(settings.file_storage_root).resolve()
        full = root / file_path
        if not full.exists():
            raise FileNotFoundError(f"File not found: {full}")
        return str(full)

    def _find_page_number(self, offset: int, page_offsets: list[tuple[int, int]]) -> int | None:
        """根据字符偏移量找到所在页码"""
        page_num = None
        for pn, po in page_offsets:
            if offset >= po:
                page_num = pn
            else:
                break
        return page_num

    async def get_status(self, content_id: str) -> dict:
        result = await self.db.execute(select(Content).where(Content.id == content_id))
        content = result.scalar_one_or_none()
        if content is None:
            raise ValueError(f"Content {content_id} not found")

        from sqlalchemy import func, case
        
        stats_result = await self.db.execute(
            select(
                func.count(ContentChunk.id).label("chunk_count"),
                func.count(case((ContentChunk.chunk_type == "text", 1))).label("text_chunks"),
                func.count(case((ContentChunk.chunk_type == "image", 1))).label("image_chunks"),
                func.count(case((ContentChunk.embedding.is_not(None), 1))).label("embedded_chunks"),
            ).where(ContentChunk.content_id == content.id)
        )
        stats = stats_result.one()

        return {
            "id": str(content.id),
            "processing_status": content.processing_status,
            "processing_error": content.processing_error,
            "has_text": bool(content.text_content),
            "has_embedding": content.embedding is not None,
            "chunk_count": stats.chunk_count or 0,
            "text_chunks": stats.text_chunks or 0,
            "image_chunks": stats.image_chunks or 0,
            "embedded_chunks": stats.embedded_chunks or 0,
        }

    # ── 各内容类型的分块处理 ──

    async def _process_pdf(self, content: Content) -> None:
        """PDF 处理：文字提取 → 语义分块 + 图片提取"""
        content_id = str(content.id)
        file_path = self._resolve_path(content.file_path)

        # 提取 PDF 文字
        pages = _extract_pdf_text(Path(file_path))
        if not pages:
            logger.warning(f"PDF 无可提取文字 - content_id={content_id}")
            return

        # 拼接全文并记录每页的字符偏移
        full_text_parts: list[str] = []
        page_offsets: list[tuple[int, int]] = []  # (page_number, cumulative_offset)
        offset = 0
        for page_num, page_text in pages:
            full_text_parts.append(page_text)
            page_offsets.append((page_num, offset))
            offset += len(page_text)
        full_text = "".join(full_text_parts)

        content.text_content = full_text

        # 语义分块
        from app.services.chunking import chunk_text
        chunks = await chunk_text(full_text, self.db)

        chunk_index = 0
        for tc in chunks:
            page_num = self._find_page_number(tc.start_offset, page_offsets)
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=chunk_index,
                chunk_type="text",
                chunk_text=tc.text,
                page_number=page_num,
                start_offset=tc.start_offset,
                end_offset=tc.end_offset,
            ))
            chunk_index += 1

        # 提取 PDF 内嵌图片
        storage_root = Path(settings.file_storage_root)
        images_dir = storage_root / "images" / content_id
        images = _extract_pdf_images(Path(file_path), images_dir)

        for page_num, img_filename in images:
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=chunk_index,
                chunk_type="image",
                image_path=f"images/{content_id}/{img_filename}",
                page_number=page_num,
            ))
            chunk_index += 1

        await self.db.flush()
        logger.info(
            f"PDF 分块完成 - content_id={content_id}, "
            f"text_chunks={len(chunks)}, image_chunks={len(images)}"
        )

    async def _process_image(self, content: Content) -> None:
        """图片处理：单块，chunk_type='image'"""
        self.db.add(ContentChunk(
            content_id=content.id,
            chunk_index=0,
            chunk_type="image",
            image_path=content.file_path,
        ))
        ocr_text = ""
        if content.file_path:
            try:
                image_path = Path(self._resolve_path(content.file_path))
                ocr_text = await _ocr_image(image_path, self.db)
            except Exception:
                logger.warning(f"图片 OCR 处理失败，保留图片块 - content_id={content.id}", exc_info=True)

        if ocr_text.strip():
            content.text_content = ocr_text.strip()
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=1,
                chunk_type="text",
                chunk_text=content.text_content,
                start_offset=0,
                end_offset=len(content.text_content),
                extra_meta={"source": "ocr"},
            ))

        await self.db.flush()
        logger.info(
            f"图片分块完成 - content_id={content.id}, "
            f"ocr_text={'yes' if ocr_text.strip() else 'no'}"
        )

    async def _process_doc(self, content: Content) -> None:
        """Office 文档处理：文字提取 → 语义分块"""
        content_id = str(content.id)
        file_path = self._resolve_path(content.file_path)
        p = Path(file_path)

        suffix = p.suffix.lower()
        metadata: dict = {}
        if suffix == ".docx":
            text, metadata = _extract_docx(p)
        elif suffix in (".xlsx", ".xls"):
            text, metadata = _extract_xlsx(p)
        else:
            text = ""

        if not text:
            logger.warning(f"文档无可提取文字 - content_id={content_id}")
            return

        content.text_content = text
        if metadata:
            content.extra_meta = {
                **(content.extra_meta or {}),
                "document_structure": metadata,
            }
        await self._text_chunk(content, text)

    async def _process_note(self, content: Content) -> None:
        """笔记处理：直接对 text_content 语义分块"""
        text = content.text_content or ""
        if not text.strip():
            logger.warning(f"笔记无文字内容 - content_id={content.id}")
            return
        await self._text_chunk(content, text)

    async def _process_web(self, content: Content) -> None:
        """网页处理：正文提取（或使用已有 text_content）→ 语义分块"""
        content_id = str(content.id)
        text = content.text_content or ""
        if not text.strip() and content.source_url:
            text = await _extract_web(content.source_url)
            content.text_content = text
        if not text.strip():
            logger.warning(f"网页无可提取文字 - content_id={content.id}")
            return
        await self._text_chunk(content, text)

        if content.source_url:
            screenshot_dir = Path(settings.file_storage_root) / "images" / content_id
            screenshot_path = await _capture_web_screenshot(content.source_url, screenshot_dir)
            if screenshot_path:
                result = await self.db.execute(
                    select(ContentChunk.chunk_index)
                    .where(ContentChunk.content_id == content.id)
                    .order_by(ContentChunk.chunk_index.desc())
                    .limit(1)
                )
                last_index = result.scalar_one_or_none()
                self.db.add(ContentChunk(
                    content_id=content.id,
                    chunk_index=(last_index + 1) if last_index is not None else 0,
                    chunk_type="image",
                    image_path=screenshot_path,
                    extra_meta={"source": "web_screenshot", "source_url": content.source_url},
                ))
                await self.db.flush()

    async def _process_audio(self, content: Content) -> None:
        """音频处理：转写 → 按字幕时间戳分块"""
        content_id = str(content.id)

        segments: list[dict] = []
        if content.file_path:
            file_path = self._resolve_path(content.file_path)
            segments = await _transcribe_audio(Path(file_path), self.db)

        if not segments:
            logger.warning(f"无转写数据，跳过音频分块 - content_id={content_id}")
            return

        full_text = " ".join(s.get("text", "") for s in segments)
        content.text_content = full_text

        for i, seg in enumerate(segments):
            seg_text = seg.get("text", "").strip()
            if not seg_text:
                continue
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=i,
                chunk_type="text",
                chunk_text=seg_text,
                time_start=seg.get("start"),
                time_end=seg.get("end"),
            ))

        await self.db.flush()
        logger.info(f"音频分块完成 - content_id={content_id}, chunks={len(segments)}")

    async def _process_video(self, content: Content) -> None:
        """视频处理：转写 + 关键帧截图 → 按时间戳分块"""
        content_id = str(content.id)

        segments: list[dict] = []
        screenshots: list[str] = []

        if content.file_path:
            # Call the module-level implementation to keep tests monkeypatchable.
            import sys
            module = sys.modules[__name__]
            file_path = self._resolve_path(content.file_path)
            screenshots_dir = Path(settings.file_storage_root) / "images" / content_id
            segments, screenshots = await module._process_video(
                Path(file_path),
                self.db,
                screenshots_dir,
            )

        if not segments:
            logger.warning(f"无转写数据，跳过视频分块 - content_id={content_id}")
            return

        full_text = " ".join(s.get("text", "") for s in segments)
        content.text_content = full_text

        chunk_index = 0
        for seg in segments:
            seg_text = seg.get("text", "").strip()
            if not seg_text:
                continue
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=chunk_index,
                chunk_type="text",
                chunk_text=seg_text,
                time_start=seg.get("start"),
                time_end=seg.get("end"),
            ))
            chunk_index += 1

        # 关键帧截图作为 image chunks
        for img_path in screenshots:
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=chunk_index,
                chunk_type="image",
                image_path=img_path,
            ))
            chunk_index += 1

        await self.db.flush()
        logger.info(
            f"视频分块完成 - content_id={content_id}, "
            f"text_chunks={len(segments)}, screenshots={len(screenshots)}"
        )

    async def _text_chunk(self, content: Content, text: str) -> None:
        """通用文本分块：语义分块 → 写入 ContentChunk（doc/note/web 共用）"""
        from app.services.chunking import chunk_text
        chunks = await chunk_text(text, self.db)

        for i, tc in enumerate(chunks):
            self.db.add(ContentChunk(
                content_id=content.id,
                chunk_index=i,
                chunk_type="text",
                chunk_text=tc.text,
                start_offset=tc.start_offset,
                end_offset=tc.end_offset,
            ))

        await self.db.flush()
        logger.info(f"文本分块完成 - content_id={content.id}, chunks={len(chunks)}")

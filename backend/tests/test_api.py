"""核心服务 & API 集成测试
用法: cd backend && venv/Scripts/python -m pytest tests/ -v
注意: Windows 上 asyncpg 连接池有限，建议按 test 逐个运行：
      pytest tests/test_api.py -k "health or storage or upload or backup" -v
"""

import uuid
from datetime import datetime, timedelta, timezone

import pytest
from httpx import AsyncClient, ASGITransport
from sqlalchemy import select

from app.main import app


# ── Fixtures ──

@pytest.fixture
async def client():
    """每个测试独立 client"""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


# ── Health ──

@pytest.mark.asyncio
async def test_health(client: AsyncClient):
    resp = await client.get("/api/health")
    assert resp.status_code == 200
    assert resp.json()["status"] == "ok"


# ── Storage ──

@pytest.mark.asyncio
async def test_storage_config(client: AsyncClient):
    resp = await client.get("/api/storage/config")
    assert resp.status_code == 200
    data = resp.json()
    assert "storage_root" in data


@pytest.mark.asyncio
async def test_storage_migrate_copies_files_and_keeps_relative_paths(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentChunk

    settings = get_settings()
    old_root = tmp_path / "old-storage"
    new_root = tmp_path / "new-storage"
    upload_path = old_root / "uploads" / "2026-06-07" / "stored.txt"
    image_path = old_root / "images" / "web_capture.png"
    upload_path.parent.mkdir(parents=True)
    image_path.parent.mkdir(parents=True)
    upload_path.write_text("stored file", encoding="utf-8")
    image_path.write_bytes(b"image-bytes")

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(old_root)
    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-storage-migrate-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="upload",
                file_path="uploads/2026-06-07/stored.txt",
                processing_status="completed",
            )
            db.add(content)
            await db.flush()
            db.add(
                ContentChunk(
                    content_id=content.id,
                    chunk_index=0,
                    chunk_type="image",
                    image_path="images/web_capture.png",
                )
            )
            await db.commit()
            content_id = content.id

        resp = await client.post("/api/storage/migrate", data={"path": str(new_root), "old_path": str(old_root)})
        assert resp.status_code == 200
        data = resp.json()
        assert data["copied"] == 2
        assert data["missing"] >= 0
        assert data["storage_root"] == str(new_root.resolve())

        assert upload_path.exists()
        assert image_path.exists()
        assert (new_root / "uploads" / "2026-06-07" / "stored.txt").read_text(encoding="utf-8") == "stored file"
        assert (new_root / "images" / "web_capture.png").read_bytes() == b"image-bytes"

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == content_id))
            refreshed = result.scalar_one()
            assert refreshed.file_path == "uploads/2026-06-07/stored.txt"
    finally:
        settings.file_storage_root = old_storage_root


# ── File Upload ──

@pytest.mark.asyncio
async def test_upload_txt(client: AsyncClient):
    resp = await client.post(
        "/api/files/upload",
        files={"file": ("test.txt", b"Hello Moyuan Test", "text/plain")},
    )
    assert resp.status_code == 201
    data = resp.json()
    assert "content_id" in data
    assert data["content_type"] == "doc"


# ── File Duplicate Check ──

@pytest.mark.asyncio
async def test_upload_duplicate_check(client: AsyncClient):
    content = b"Duplicate Test Unique Content"
    resp1 = await client.post(
        "/api/files/upload",
        files={"file": ("dup.txt", content, "text/plain")},
    )
    assert resp1.status_code == 201

    resp2 = await client.post(
        "/api/files/check-duplicate",
        files={"file": ("dup2.txt", content, "text/plain")},
    )
    assert resp2.status_code == 200
    data = resp2.json()
    assert data["is_duplicate"] is True
    assert "md5" in data["duplicates"][0]["match_types"]


@pytest.mark.asyncio
async def test_upload_duplicate_check_matches_filename(client: AsyncClient):
    resp1 = await client.post(
        "/api/files/upload",
        files={"file": ("same-name.txt", b"first content", "text/plain")},
    )
    assert resp1.status_code == 201

    resp2 = await client.post(
        "/api/files/check-duplicate",
        files={"file": ("same-name.txt", b"different content", "text/plain")},
    )
    assert resp2.status_code == 200
    data = resp2.json()
    assert data["is_duplicate"] is True
    assert any("filename" in d["match_types"] for d in data["duplicates"])


@pytest.mark.asyncio
async def test_upload_duplicate_check_is_brain_scoped(client: AsyncClient):
    brain_a_resp = await client.post("/api/brains", json={"name": f"dup-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"dup-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]
    content = f"scoped duplicate {uuid.uuid4().hex}".encode("utf-8")

    upload_resp = await client.post(
        "/api/files/upload",
        data={"brain_id": brain_a},
        files={"file": ("scoped-dup.txt", content, "text/plain")},
    )
    assert upload_resp.status_code == 201

    same_brain_resp = await client.post(
        "/api/files/check-duplicate",
        data={"brain_id": brain_a},
        files={"file": ("scoped-dup-copy.txt", content, "text/plain")},
    )
    other_brain_resp = await client.post(
        "/api/files/check-duplicate",
        data={"brain_id": brain_b},
        files={"file": ("scoped-dup-copy.txt", content, "text/plain")},
    )

    assert same_brain_resp.status_code == 200
    assert same_brain_resp.json()["is_duplicate"] is True
    assert other_brain_resp.status_code == 200
    assert other_brain_resp.json()["is_duplicate"] is False


@pytest.mark.asyncio
async def test_upload_overwrite_rejects_cross_brain_target(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Category, Collection, CollectionItem, Content, ContentCategory

    brain_a_resp = await client.post("/api/brains", json={"name": f"overwrite-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"overwrite-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]

    original_resp = await client.post(
        "/api/files/upload",
        data={"brain_id": brain_b},
        files={"file": ("overwrite-target.txt", b"original", "text/plain")},
    )
    assert original_resp.status_code == 201
    original_id = original_resp.json()["content_id"]

    cross_resp = await client.post(
        "/api/files/upload",
        data={"brain_id": brain_a, "overwrite_content_id": original_id},
        files={"file": ("overwrite-new.txt", b"new", "text/plain")},
    )
    missing_resp = await client.post(
        "/api/files/upload",
        data={"brain_id": brain_a, "overwrite_content_id": str(uuid.uuid4())},
        files={"file": ("overwrite-missing.txt", b"new", "text/plain")},
    )

    assert cross_resp.status_code == 400
    assert missing_resp.status_code == 404
    async with async_session_factory() as db:
        original = await db.get(Content, uuid.UUID(original_id))
        assert original is not None
        assert original.is_deleted is False


@pytest.mark.asyncio
async def test_folder_upload_preserves_relative_path(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Category, Collection, CollectionItem, Content, ContentCategory
    import app.services.file as file_service

    settings = get_settings()
    old_root = settings.file_storage_root
    old_service_root = file_service.settings.file_storage_root
    storage_root = tmp_path / "storage"

    try:
        settings.file_storage_root = str(storage_root)
        file_service.settings.file_storage_root = str(storage_root)
        file_bytes = f"folder lesson {uuid.uuid4()}".encode("utf-8")
        import_root = f"Photography Course {uuid.uuid4().hex[:8]}"
        import_relative_path = f"{import_root}/Week 1/lesson.txt"

        resp = await client.post(
            "/api/files/upload",
            data={
                "import_relative_path": import_relative_path,
                "import_batch_id": "batch-folder-test",
            },
            files={"file": ("lesson.txt", file_bytes, "text/plain")},
        )
        bad_resp = await client.post(
            "/api/files/upload",
            data={"import_relative_path": "../escape.txt", "import_batch_id": "batch-folder-test"},
            files={"file": ("escape.txt", b"nope", "text/plain")},
        )

        assert resp.status_code == 201
        assert bad_resp.status_code == 400
        data = resp.json()
        stored_path = data["file_path"].replace("\\", "/")
        assert data["title"] == "lesson"
        assert stored_path.startswith("imports/")
        assert stored_path.endswith(import_relative_path)
        assert (storage_root / data["file_path"]).read_bytes() == file_bytes

        async with async_session_factory() as db:
            content = await db.get(Content, uuid.UUID(data["content_id"]))
            assert content.extra_meta["original_filename"] == "lesson.txt"
            assert content.extra_meta["import_relative_path"] == import_relative_path
            assert content.extra_meta["import_root"] == import_root
            assert content.extra_meta["import_batch_id"] == "batch-folder-test"
            collection_result = await db.execute(
                select(Collection).where(Collection.name == import_root, Collection.brain_id.is_(None))
            )
            collection = collection_result.scalar_one()
            item_result = await db.execute(
                select(CollectionItem).where(
                    CollectionItem.collection_id == collection.id,
                    CollectionItem.content_id == content.id,
                )
            )
            assert item_result.scalar_one_or_none() is not None
            assert content.extra_meta["import_collection_id"] == str(collection.id)
            root_category_result = await db.execute(
                select(Category).where(Category.name == import_root, Category.parent_id.is_(None), Category.brain_id.is_(None))
            )
            root_category = root_category_result.scalar_one()
            week_category_result = await db.execute(
                select(Category).where(Category.name == "Week 1", Category.parent_id == root_category.id)
            )
            week_category = week_category_result.scalar_one()
            category_link_result = await db.execute(
                select(ContentCategory).where(
                    ContentCategory.content_id == content.id,
                    ContentCategory.category_id == week_category.id,
                )
            )
            assert category_link_result.scalar_one_or_none() is not None
            assert content.extra_meta["import_category_id"] == str(week_category.id)
            import_category_id = content.extra_meta["import_category_id"]
            study_started_at = "2026-01-01T00:00:00+00:00"
            study_completed_at = datetime.now(timezone.utc).isoformat()
            content.extra_meta = {
                **content.extra_meta,
                "study_status": "completed",
                "study_started_at": study_started_at,
                "study_completed_at": study_completed_at,
            }
            deleted_content = Content(
                title=f"deleted-import-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="manual",
                is_deleted=True,
                deleted_at=datetime.now(timezone.utc),
            )
            db.add(deleted_content)
            await db.flush()
            db.add(CollectionItem(collection_id=collection.id, content_id=deleted_content.id, sort_order=99))
            await db.commit()
            collection_id = str(collection.id)
            content_id = str(content.id)
            deleted_content_id = str(deleted_content.id)

        detail_resp = await client.get(f"/api/collections/{collection_id}")
        assert detail_resp.status_code == 200
        detail_items = detail_resp.json()["items"]
        detail_ids = [item["content_id"] for item in detail_items]
        assert content_id in detail_ids
        assert deleted_content_id not in detail_ids
        detail_item = next(item for item in detail_items if item["content_id"] == content_id)
        assert detail_item["import_relative_path"] == import_relative_path
        assert detail_item["folder_path"] == f"{import_root}/Week 1"
        assert detail_item["import_root"] == import_root
        assert detail_item["import_category_id"] == import_category_id
        assert detail_item["study_status"] == "completed"
        assert detail_item["study_started_at"] == study_started_at
        assert detail_item["study_completed_at"] == study_completed_at

        list_resp = await client.get("/api/collections?page=1&page_size=100")
        assert list_resp.status_code == 200
        listed_collection = next(item for item in list_resp.json()["items"] if item["id"] == collection_id)
        assert listed_collection["item_count"] == 1
        assert listed_collection["completed_count"] == 1
        assert listed_collection["in_progress_count"] == 0
        assert listed_collection["progress_percent"] == 100
        assert listed_collection["resume_content_id"] is None
    finally:
        settings.file_storage_root = old_root
        file_service.settings.file_storage_root = old_service_root


@pytest.mark.asyncio
async def test_collection_list_returns_resume_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Collection, CollectionItem, Content

    completed_title = f"collection-resume-completed-{uuid.uuid4().hex[:8]}"
    progress_title = f"collection-resume-progress-{uuid.uuid4().hex[:8]}"
    not_started_title = f"collection-resume-unset-{uuid.uuid4().hex[:8]}"
    collection_name = f"collection-resume-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        collection = Collection(name=collection_name)
        completed = Content(
            title=completed_title,
            content_type="doc",
            source_type="manual",
            extra_meta={"study_status": "completed"},
        )
        in_progress = Content(
            title=progress_title,
            content_type="doc",
            source_type="manual",
            extra_meta={"study_status": "in_progress"},
        )
        not_started = Content(
            title=not_started_title,
            content_type="doc",
            source_type="manual",
        )
        db.add_all([collection, completed, in_progress, not_started])
        await db.flush()
        db.add_all([
            CollectionItem(collection_id=collection.id, content_id=completed.id, sort_order=1),
            CollectionItem(collection_id=collection.id, content_id=not_started.id, sort_order=2),
            CollectionItem(collection_id=collection.id, content_id=in_progress.id, sort_order=3),
        ])
        await db.commit()
        collection_id = str(collection.id)
        in_progress_id = str(in_progress.id)

    list_resp = await client.get("/api/collections?page=1&page_size=100")
    detail_resp = await client.get(f"/api/collections/{collection_id}")

    assert list_resp.status_code == 200
    listed_collection = next(item for item in list_resp.json()["items"] if item["id"] == collection_id)
    assert listed_collection["resume_content_id"] == in_progress_id
    assert listed_collection["resume_content_title"] == progress_title
    assert listed_collection["resume_study_status"] == "in_progress"

    assert detail_resp.status_code == 200
    assert detail_resp.json()["collection"]["resume_content_id"] == in_progress_id


@pytest.mark.asyncio
async def test_collection_list_filters_by_query_and_progress(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Collection, CollectionItem, Content

    prefix = f"collection-filter-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        completed_col = Collection(name=f"{prefix}-completed", description="photography course")
        progress_col = Collection(name=f"{prefix}-progress", description="ai learning")
        not_done_col = Collection(name=f"{prefix}-not-done", description="mixed study")
        empty_col = Collection(name=f"{prefix}-empty", description="empty course")
        other_col = Collection(name=f"unrelated-{uuid.uuid4().hex[:8]}", description="other")
        completed_content = Content(
            title=f"{prefix}-completed-content",
            content_type="doc",
            source_type="manual",
            extra_meta={"study_status": "completed"},
        )
        progress_content = Content(
            title=f"{prefix}-progress-content",
            content_type="doc",
            source_type="manual",
            extra_meta={"study_status": "in_progress"},
        )
        not_done_completed = Content(
            title=f"{prefix}-not-done-completed",
            content_type="doc",
            source_type="manual",
            extra_meta={"study_status": "completed"},
        )
        not_done_unset = Content(
            title=f"{prefix}-not-done-unset",
            content_type="doc",
            source_type="manual",
        )
        db.add_all([
            completed_col,
            progress_col,
            not_done_col,
            empty_col,
            other_col,
            completed_content,
            progress_content,
            not_done_completed,
            not_done_unset,
        ])
        await db.flush()
        db.add_all([
            CollectionItem(collection_id=completed_col.id, content_id=completed_content.id, sort_order=1),
            CollectionItem(collection_id=progress_col.id, content_id=progress_content.id, sort_order=1),
            CollectionItem(collection_id=not_done_col.id, content_id=not_done_completed.id, sort_order=1),
            CollectionItem(collection_id=not_done_col.id, content_id=not_done_unset.id, sort_order=2),
        ])
        await db.commit()

    query_resp = await client.get(f"/api/collections?page=1&page_size=100&q={prefix}")
    query_page_2_resp = await client.get(f"/api/collections?page=2&page_size=2&q={prefix}")
    completed_resp = await client.get(f"/api/collections?page=1&page_size=100&q={prefix}&progress=completed")
    progress_resp = await client.get(f"/api/collections?page=1&page_size=100&q={prefix}&progress=in_progress")
    not_done_resp = await client.get(f"/api/collections?page=1&page_size=100&q={prefix}&progress=not_done")

    assert query_resp.status_code == 200
    query_body = query_resp.json()
    assert query_body["total"] == 4
    assert query_body["page"] == 1
    assert query_body["page_size"] == 100
    query_names = {item["name"] for item in query_body["items"]}
    assert query_names == {
        f"{prefix}-completed",
        f"{prefix}-progress",
        f"{prefix}-not-done",
        f"{prefix}-empty",
    }
    assert query_page_2_resp.status_code == 200
    query_page_2_body = query_page_2_resp.json()
    assert query_page_2_body["total"] == 4
    assert query_page_2_body["page"] == 2
    assert query_page_2_body["page_size"] == 2
    assert len(query_page_2_body["items"]) == 2

    assert completed_resp.status_code == 200
    assert completed_resp.json()["total"] == 1
    assert [item["name"] for item in completed_resp.json()["items"]] == [f"{prefix}-completed"]

    assert progress_resp.status_code == 200
    assert progress_resp.json()["total"] == 1
    assert [item["name"] for item in progress_resp.json()["items"]] == [f"{prefix}-progress"]

    assert not_done_resp.status_code == 200
    assert not_done_resp.json()["total"] == 2
    not_done_names = {item["name"] for item in not_done_resp.json()["items"]}
    assert not_done_names == {f"{prefix}-progress", f"{prefix}-not-done"}


@pytest.mark.asyncio
async def test_folder_duplicate_check_uses_relative_path(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    import app.services.file as file_service

    settings = get_settings()
    old_root = settings.file_storage_root
    old_service_root = file_service.settings.file_storage_root
    storage_root = tmp_path / "storage"
    root = f"Course {uuid.uuid4().hex[:8]}"

    try:
        settings.file_storage_root = str(storage_root)
        file_service.settings.file_storage_root = str(storage_root)

        upload_resp = await client.post(
            "/api/files/upload",
            data={
                "import_relative_path": f"{root}/Week 1/lesson.txt",
                "import_batch_id": "dup-folder-test",
            },
            files={"file": ("lesson.txt", b"week one", "text/plain")},
        )
        different_path_resp = await client.post(
            "/api/files/check-duplicate",
            data={"import_relative_path": f"{root}/Week 2/lesson.txt"},
            files={"file": ("lesson.txt", b"week two", "text/plain")},
        )
        same_path_resp = await client.post(
            "/api/files/check-duplicate",
            data={"import_relative_path": f"{root}/Week 1/lesson.txt"},
            files={"file": ("lesson.txt", b"updated week one", "text/plain")},
        )
        bad_path_resp = await client.post(
            "/api/files/check-duplicate",
            data={"import_relative_path": "../escape.txt"},
            files={"file": ("escape.txt", b"nope", "text/plain")},
        )

        assert upload_resp.status_code == 201
        assert different_path_resp.status_code == 200
        assert different_path_resp.json()["is_duplicate"] is False
        assert same_path_resp.status_code == 200
        same_path = same_path_resp.json()
        assert same_path["is_duplicate"] is True
        assert "relative_path" in same_path["duplicates"][0]["match_types"]
        assert bad_path_resp.status_code == 400
    finally:
        settings.file_storage_root = old_root
        file_service.settings.file_storage_root = old_service_root


# ── File List ──

@pytest.mark.asyncio
async def test_file_list(client: AsyncClient):
    resp = await client.get("/api/files?page=1&page_size=5")
    assert resp.status_code == 200
    data = resp.json()
    assert "items" in data
    assert "total" in data
    assert isinstance(data["items"], list)


@pytest.mark.asyncio
async def test_file_list_filters_by_processing_status(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content

    failed_title = f"pytest-failed-{uuid.uuid4().hex[:8]}"
    completed_title = f"pytest-completed-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        db.add_all([
            Content(
                title=failed_title,
                content_type="doc",
                source_type="manual",
                text_content="failed content",
                processing_status="failed",
            ),
            Content(
                title=completed_title,
                content_type="doc",
                source_type="manual",
                text_content="completed content",
                processing_status="completed",
            ),
        ])
        await db.commit()

    resp = await client.get("/api/files?processing_status=failed&page=1&page_size=100")
    assert resp.status_code == 200
    titles = {item["title"] for item in resp.json()["items"]}
    assert failed_title in titles
    assert completed_title not in titles


@pytest.mark.asyncio
async def test_file_list_filters_by_study_status(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content

    completed_title = f"pytest-study-completed-{uuid.uuid4().hex[:8]}"
    in_progress_title = f"pytest-study-progress-{uuid.uuid4().hex[:8]}"
    implicit_not_started_title = f"pytest-study-unset-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        db.add_all([
            Content(
                title=completed_title,
                content_type="doc",
                source_type="manual",
                text_content="completed study content",
                extra_meta={"study_status": "completed"},
            ),
            Content(
                title=in_progress_title,
                content_type="doc",
                source_type="manual",
                text_content="in progress study content",
                extra_meta={"study_status": "in_progress"},
            ),
            Content(
                title=implicit_not_started_title,
                content_type="doc",
                source_type="manual",
                text_content="not started study content",
                extra_meta={"import_relative_path": "Course/lesson.txt"},
            ),
        ])
        await db.commit()

    completed_resp = await client.get("/api/files?study_status=completed&page=1&page_size=100")
    not_started_resp = await client.get("/api/files?study_status=not_started&page=1&page_size=100")
    invalid_resp = await client.get("/api/files?study_status=paused&page=1&page_size=100")

    assert completed_resp.status_code == 200
    completed_titles = {item["title"] for item in completed_resp.json()["items"]}
    assert completed_title in completed_titles
    assert in_progress_title not in completed_titles

    assert not_started_resp.status_code == 200
    not_started_titles = {item["title"] for item in not_started_resp.json()["items"]}
    assert implicit_not_started_title in not_started_titles
    assert completed_title not in not_started_titles
    assert invalid_resp.status_code == 400


@pytest.mark.asyncio
async def test_file_list_filters_by_query_and_brain(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, Content

    keyword = f"needle-{uuid.uuid4().hex[:8]}"
    brain_a = Brain(name=f"query-brain-a-{uuid.uuid4().hex[:8]}")
    brain_b = Brain(name=f"query-brain-b-{uuid.uuid4().hex[:8]}")
    async with async_session_factory() as db:
        db.add_all([brain_a, brain_b])
        await db.flush()
        db.add_all([
            Content(
                title=f"{keyword} title",
                content_type="doc",
                source_type="manual",
                text_content="ordinary body",
                brain_id=brain_a.id,
            ),
            Content(
                title="body match",
                content_type="doc",
                source_type="manual",
                text_content=f"body has {keyword}",
                brain_id=brain_a.id,
            ),
            Content(
                title=f"{keyword} wrong brain",
                content_type="doc",
                source_type="manual",
                text_content="ordinary body",
                brain_id=brain_b.id,
            ),
        ])
        await db.commit()
        brain_a_id = str(brain_a.id)

    resp = await client.get(f"/api/files?q={keyword}&brain_id={brain_a_id}&page=1&page_size=100")
    assert resp.status_code == 200
    titles = {item["title"] for item in resp.json()["items"]}
    assert f"{keyword} title" in titles
    assert "body match" in titles
    assert f"{keyword} wrong brain" not in titles


# ── Tags ──

@pytest.mark.asyncio
async def test_recycle_cleanup_removes_items_older_than_30_days(client: AsyncClient, tmp_path):
    from datetime import datetime, timedelta, timezone
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    import app.services.file as file_service

    settings = get_settings()
    old_root = settings.file_storage_root
    old_service_root = file_service.settings.file_storage_root
    storage_root = tmp_path / "storage"
    storage_root.mkdir()
    rel_path = "old/deleted.txt"
    file_path = storage_root / rel_path
    file_path.parent.mkdir(parents=True)
    file_path.write_text("expired deleted file", encoding="utf-8")

    try:
        settings.file_storage_root = str(storage_root)
        file_service.settings.file_storage_root = str(storage_root)

        expired_id = uuid.uuid4()
        active_id = uuid.uuid4()
        async with async_session_factory() as db:
            db.add_all([
                Content(
                    id=expired_id,
                    title=f"expired-{uuid.uuid4().hex[:8]}",
                    content_type="doc",
                    source_type="upload",
                    file_path=rel_path,
                    is_deleted=True,
                    deleted_at=datetime.now(timezone.utc) - timedelta(days=31),
                ),
                Content(
                    id=active_id,
                    title=f"recent-{uuid.uuid4().hex[:8]}",
                    content_type="doc",
                    source_type="upload",
                    is_deleted=True,
                    deleted_at=datetime.now(timezone.utc) - timedelta(days=3),
                ),
            ])
            await db.commit()

        resp = await client.post("/api/recycle/cleanup")
        assert resp.status_code == 200
        data = resp.json()
        assert data["deleted_count"] >= 1
        assert data["retention_days"] == 30
        assert not file_path.exists()

        async with async_session_factory() as db:
            expired = await db.get(Content, expired_id)
            active = await db.get(Content, active_id)
            assert expired is None
            assert active is not None
    finally:
        settings.file_storage_root = old_root
        file_service.settings.file_storage_root = old_service_root


@pytest.mark.asyncio
async def test_recycle_batch_actions_are_brain_scoped_and_require_deleted_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, Content

    brain_a = uuid.uuid4()
    brain_b = uuid.uuid4()
    active_id = uuid.uuid4()
    deleted_a_id = uuid.uuid4()
    deleted_b_id = uuid.uuid4()
    permanent_a_id = uuid.uuid4()

    async with async_session_factory() as db:
        db.add_all([
            Brain(id=brain_a, name=f"recycle-a-{uuid.uuid4().hex[:8]}"),
            Brain(id=brain_b, name=f"recycle-b-{uuid.uuid4().hex[:8]}"),
            Content(
                id=active_id,
                title=f"active-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="manual",
                brain_id=brain_a,
                is_deleted=False,
            ),
            Content(
                id=deleted_a_id,
                title=f"deleted-a-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="manual",
                brain_id=brain_a,
                is_deleted=True,
                deleted_at=datetime.now(timezone.utc),
            ),
            Content(
                id=deleted_b_id,
                title=f"deleted-b-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="manual",
                brain_id=brain_b,
                is_deleted=True,
                deleted_at=datetime.now(timezone.utc),
            ),
            Content(
                id=permanent_a_id,
                title=f"permanent-a-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="manual",
                brain_id=brain_a,
                is_deleted=True,
                deleted_at=datetime.now(timezone.utc),
            ),
        ])
        await db.commit()

    active_permanent_resp = await client.delete(f"/api/recycle/{active_id}/permanent")
    restore_resp = await client.post(
        "/api/contents/batch",
        json={
            "ids": [str(deleted_a_id), str(deleted_b_id)],
            "action": "restore",
            "brain_id": str(brain_a),
        },
    )
    delete_resp = await client.post(
        "/api/contents/batch",
        json={
            "ids": [str(permanent_a_id), str(deleted_b_id)],
            "action": "permanent_delete",
            "brain_id": str(brain_a),
        },
    )
    soft_delete_resp = await client.post(
        "/api/contents/batch",
        json={"ids": [str(active_id)], "action": "delete", "brain_id": str(brain_a)},
    )

    assert active_permanent_resp.status_code == 404
    assert restore_resp.status_code == 200
    assert restore_resp.json()["updated"] == 1
    assert delete_resp.status_code == 200
    assert delete_resp.json()["updated"] == 1
    assert soft_delete_resp.status_code == 200
    assert soft_delete_resp.json()["updated"] == 1

    async with async_session_factory() as db:
        active = await db.get(Content, active_id)
        restored_a = await db.get(Content, deleted_a_id)
        still_deleted_b = await db.get(Content, deleted_b_id)
        removed_a = await db.get(Content, permanent_a_id)
        assert active.is_deleted is True
        assert active.deleted_at is not None
        assert restored_a.is_deleted is False
        assert restored_a.deleted_at is None
        assert still_deleted_b.is_deleted is True
        assert removed_a is None


@pytest.mark.asyncio
async def test_tag_crud(client: AsyncClient):
    tag_name = f"pytest-tag-{uuid.uuid4().hex[:8]}"
    resp = await client.post("/api/tags", json={"name": tag_name, "color": "#ff0000"})
    assert resp.status_code == 201

    resp2 = await client.get("/api/tags?page=1&page_size=100")
    assert resp2.status_code == 200


# ── Categories ──

@pytest.mark.asyncio
async def test_category_crud(client: AsyncClient):
    resp = await client.post("/api/categories", json={"name": "pytest-cat"})
    assert resp.status_code == 201
    cid = resp.json()["id"]

    resp2 = await client.get("/api/categories/tree")
    assert resp2.status_code == 200

    resp3 = await client.delete(f"/api/categories/{cid}")
    assert resp3.status_code == 200


@pytest.mark.asyncio
async def test_brain_scoped_tags_categories_and_collections(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content

    brain_a_resp = await client.post("/api/brains", json={"name": f"org-brain-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"org-brain-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = uuid.UUID(brain_a_resp.json()["id"])
    brain_b = uuid.UUID(brain_b_resp.json()["id"])

    tag_a = await client.post("/api/tags", json={"name": f"tag-a-{uuid.uuid4().hex[:8]}", "brain_id": str(brain_a)})
    tag_b = await client.post("/api/tags", json={"name": f"tag-b-{uuid.uuid4().hex[:8]}", "brain_id": str(brain_b)})
    assert tag_a.status_code == 201
    assert tag_b.status_code == 201

    tags_a = await client.get(f"/api/tags?brain_id={brain_a}")
    assert tags_a.status_code == 200
    tag_ids_a = {item["id"] for item in tags_a.json()}
    assert tag_a.json()["id"] in tag_ids_a
    assert tag_b.json()["id"] not in tag_ids_a

    cat_a = await client.post("/api/categories", json={"name": f"cat-a-{uuid.uuid4().hex[:8]}", "brain_id": str(brain_a)})
    cat_b = await client.post("/api/categories", json={"name": f"cat-b-{uuid.uuid4().hex[:8]}", "brain_id": str(brain_b)})
    assert cat_a.status_code == 201
    assert cat_b.status_code == 201

    cats_a = await client.get(f"/api/categories?brain_id={brain_a}")
    assert cats_a.status_code == 200
    cat_ids_a = {item["id"] for item in cats_a.json()}
    assert cat_a.json()["id"] in cat_ids_a
    assert cat_b.json()["id"] not in cat_ids_a

    collection_a = await client.post(
        "/api/collections",
        json={"name": f"collection-a-{uuid.uuid4().hex[:8]}", "brain_id": str(brain_a)},
    )
    assert collection_a.status_code == 201
    collections_a = await client.get(f"/api/collections?brain_id={brain_a}")
    assert collections_a.status_code == 200
    assert collection_a.json()["id"] in {item["id"] for item in collections_a.json()["items"]}

    async with async_session_factory() as db:
        content_b = Content(
            title=f"pytest-brain-b-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="brain b content",
            brain_id=brain_b,
        )
        db.add(content_b)
        await db.commit()
        content_b_id = str(content_b.id)

    add_cross = await client.post(f"/api/collections/{collection_a.json()['id']}/add", json={"content_id": content_b_id})
    assert add_cross.status_code == 400


@pytest.mark.asyncio
async def test_organization_create_rejects_unknown_brain_id(client: AsyncClient):
    missing_brain = str(uuid.uuid4())

    tag_resp = await client.post("/api/tags", json={"name": f"unknown-tag-{uuid.uuid4().hex[:8]}", "brain_id": missing_brain})
    category_resp = await client.post("/api/categories", json={"name": f"unknown-cat-{uuid.uuid4().hex[:8]}", "brain_id": missing_brain})
    collection_resp = await client.post("/api/collections", json={"name": f"unknown-col-{uuid.uuid4().hex[:8]}", "brain_id": missing_brain})

    assert tag_resp.status_code == 404
    assert category_resp.status_code == 404
    assert collection_resp.status_code == 404


@pytest.mark.asyncio
async def test_organization_reads_reject_unknown_brain_id(client: AsyncClient):
    missing_brain = str(uuid.uuid4())

    tags_resp = await client.get(f"/api/tags?brain_id={missing_brain}")
    categories_resp = await client.get(f"/api/categories?brain_id={missing_brain}")
    category_tree_resp = await client.get(f"/api/categories/tree?brain_id={missing_brain}")
    collections_resp = await client.get(f"/api/collections?brain_id={missing_brain}")
    favorites_resp = await client.get(f"/api/collections/favorites?brain_id={missing_brain}")

    assert tags_resp.status_code == 404
    assert categories_resp.status_code == 404
    assert category_tree_resp.status_code == 404
    assert collections_resp.status_code == 404
    assert favorites_resp.status_code == 404


@pytest.mark.asyncio
async def test_invalid_brain_id_returns_400_for_scoped_endpoints(client: AsyncClient):
    bad_brain = "not-a-uuid"

    responses = [
        await client.get(f"/api/tags?brain_id={bad_brain}"),
        await client.post("/api/tags", json={"name": f"bad-tag-{uuid.uuid4().hex[:8]}", "brain_id": bad_brain}),
        await client.get(f"/api/categories?brain_id={bad_brain}"),
        await client.post("/api/categories", json={"name": f"bad-cat-{uuid.uuid4().hex[:8]}", "brain_id": bad_brain}),
        await client.get(f"/api/collections?brain_id={bad_brain}"),
        await client.post("/api/collections", json={"name": f"bad-col-{uuid.uuid4().hex[:8]}", "brain_id": bad_brain}),
        await client.get(f"/api/files?brain_id={bad_brain}"),
        await client.get(f"/api/storage/stats?brain_id={bad_brain}"),
        await client.get(f"/api/notes?brain_id={bad_brain}"),
        await client.post("/api/notes", json={"title": f"bad-note-{uuid.uuid4().hex[:8]}", "brain_id": bad_brain}),
        await client.post("/api/contents/batch", json={"ids": [], "action": "star", "brain_id": bad_brain}),
    ]

    assert all(resp.status_code == 400 for resp in responses)


@pytest.mark.asyncio
async def test_notes_are_brain_scoped(client: AsyncClient):
    brain_a_resp = await client.post("/api/brains", json={"name": f"notes-brain-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"notes-brain-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]

    note_a = await client.post(
        "/api/notes",
        json={"title": f"note-a-{uuid.uuid4().hex[:8]}", "content": "a", "brain_id": str(brain_a)},
    )
    note_b = await client.post(
        "/api/notes",
        json={"title": f"note-b-{uuid.uuid4().hex[:8]}", "content": "b", "brain_id": str(brain_b)},
    )
    assert note_a.status_code == 200
    assert note_b.status_code == 200

    notes_a = await client.get(f"/api/notes?brain_id={brain_a}&page_size=100")
    assert notes_a.status_code == 200
    ids_a = {item["id"] for item in notes_a.json()["items"]}
    assert note_a.json()["id"] in ids_a
    assert note_b.json()["id"] not in ids_a


@pytest.mark.asyncio
async def test_notes_and_prompt_templates_reject_unknown_brain_id(client: AsyncClient):
    missing_brain = str(uuid.uuid4())

    note_resp = await client.post(
        "/api/notes",
        json={"title": f"unknown-note-{uuid.uuid4().hex[:8]}", "content": "x", "brain_id": missing_brain},
    )
    note_list_resp = await client.get(f"/api/notes?brain_id={missing_brain}")
    quiz_template_resp = await client.put(
        "/api/ai/quiz-template",
        json={"system_prompt": "x", "user_prompt_template": "y", "brain_id": missing_brain},
    )
    qa_template_resp = await client.post(f"/api/ai/qa-template/reset?brain_id={missing_brain}")

    assert note_resp.status_code == 404
    assert note_list_resp.status_code == 404
    assert quiz_template_resp.status_code == 404
    assert qa_template_resp.status_code == 404


@pytest.mark.asyncio
async def test_ai_quiz_and_ask_reject_unknown_brain_id(client: AsyncClient):
    missing_brain = str(uuid.uuid4())

    quiz_resp = await client.post(
        "/api/ai/quiz",
        json={"content_ids": [], "brain_id": missing_brain},
    )
    wrong_quiz_resp = await client.post(
        "/api/ai/wrong_quiz",
        json={"wrong_question_texts": ["错题"], "brain_id": missing_brain},
    )
    ask_resp = await client.post(
        "/api/ai/ask",
        json={"question": "这个知识库有什么？", "brain_id": missing_brain},
    )

    assert quiz_resp.status_code == 404
    assert wrong_quiz_resp.status_code == 404
    assert ask_resp.status_code == 404


@pytest.mark.asyncio
async def test_wrong_answers_are_paginated_scoped_and_reappear_after_new_mistake(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Collection, CollectionItem, Content, Question, QuestionRecord

    prefix = f"wrong-page-{uuid.uuid4().hex[:8]}"
    base_time = datetime(2026, 1, 1, tzinfo=timezone.utc)

    async with async_session_factory() as db:
        content = Content(
            title=f"{prefix}-content",
            content_type="doc",
            source_type="manual",
            text_content="wrong answer scope content",
        )
        outside_content = Content(
            title=f"{prefix}-outside",
            content_type="doc",
            source_type="manual",
            text_content="outside wrong answer content",
        )
        collection = Collection(name=f"{prefix}-collection")
        db.add_all([content, outside_content, collection])
        await db.flush()
        db.add(CollectionItem(collection_id=collection.id, content_id=content.id, sort_order=1))

        scoped_question_ids: list[str] = []
        for index in range(12):
            question = Question(
                content_id=content.id,
                q_type="single",
                question=f"{prefix}-question-{index}",
                options=["A", "B"],
                answer="A",
            )
            db.add(question)
            await db.flush()
            scoped_question_ids.append(str(question.id))
            db.add(QuestionRecord(
                question_id=question.id,
                user_answer=f"wrong-{index}",
                is_correct=False,
                answered_at=base_time + timedelta(minutes=index),
            ))

        first_question_id = scoped_question_ids[0]
        db.add(QuestionRecord(
            question_id=uuid.UUID(first_question_id),
            user_answer="latest-wrong",
            is_correct=False,
            answered_at=base_time + timedelta(minutes=100),
        ))

        outside_question = Question(
            content_id=outside_content.id,
            q_type="single",
            question=f"{prefix}-outside-question",
            options=["A", "B"],
            answer="A",
        )
        resolved_question = Question(
            content_id=content.id,
            q_type="single",
            question=f"{prefix}-resolved-question",
            options=["A", "B"],
            answer="A",
        )
        db.add_all([outside_question, resolved_question])
        await db.flush()
        db.add_all([
            QuestionRecord(
                question_id=outside_question.id,
                user_answer="outside-wrong",
                is_correct=False,
                answered_at=base_time + timedelta(minutes=120),
            ),
            QuestionRecord(
                question_id=resolved_question.id,
                user_answer="old-wrong",
                is_correct=False,
                answered_at=base_time + timedelta(minutes=130),
            ),
            QuestionRecord(
                question_id=resolved_question.id,
                user_answer="corrected",
                is_correct=True,
                answered_at=base_time + timedelta(minutes=140),
            ),
        ])
        await db.commit()
        collection_id = str(collection.id)

    page_1 = await client.get(f"/api/ai/quiz/wrong?scope_type=collection&scope_id={collection_id}&page=1&page_size=5")
    page_3 = await client.get(f"/api/ai/quiz/wrong?scope_type=collection&scope_id={collection_id}&page=3&page_size=5")

    assert page_1.status_code == 200
    page_1_body = page_1.json()
    assert page_1_body["total"] == 12
    assert page_1_body["page"] == 1
    assert page_1_body["page_size"] == 5
    assert len(page_1_body["questions"]) == 5
    assert page_1_body["questions"][0]["id"] == first_question_id
    assert page_1_body["questions"][0]["user_answer"] == "latest-wrong"

    assert page_3.status_code == 200
    assert page_3.json()["total"] == 12
    assert len(page_3.json()["questions"]) == 2

    remove_resp = await client.delete(f"/api/ai/quiz/wrong/{first_question_id}")
    after_remove = await client.get(f"/api/ai/quiz/wrong?scope_type=collection&scope_id={collection_id}&page=1&page_size=20")

    assert remove_resp.status_code == 200
    assert after_remove.status_code == 200
    assert after_remove.json()["total"] == 11
    assert first_question_id not in {item["id"] for item in after_remove.json()["questions"]}

    record_again = await client.post(
        "/api/ai/quiz/record",
        json={"question_id": first_question_id, "user_answer": "wrong-again", "is_correct": False},
    )
    after_reappear = await client.get(f"/api/ai/quiz/wrong?scope_type=collection&scope_id={collection_id}&page=1&page_size=20")

    assert record_again.status_code == 200
    assert after_reappear.status_code == 200
    assert after_reappear.json()["total"] == 12
    assert first_question_id in {item["id"] for item in after_reappear.json()["questions"]}


@pytest.mark.asyncio
async def test_quiz_history_is_paginated_and_scoped(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Collection, CollectionItem, Content, Question

    prefix = f"history-page-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        content = Content(
            title=f"{prefix}-content",
            content_type="doc",
            source_type="manual",
            text_content="quiz history scope content",
        )
        outside_content = Content(
            title=f"{prefix}-outside",
            content_type="doc",
            source_type="manual",
            text_content="outside quiz history content",
        )
        collection = Collection(name=f"{prefix}-collection")
        db.add_all([content, outside_content, collection])
        await db.flush()
        db.add(CollectionItem(collection_id=collection.id, content_id=content.id, sort_order=1))

        for index in range(7):
            db.add(Question(
                content_id=content.id,
                q_type="single",
                question=f"{prefix}-scoped-{index}",
                options=["A", "B"],
                answer="A",
                created_at=datetime(2026, 1, 1, tzinfo=timezone.utc) + timedelta(minutes=index),
            ))
        db.add(Question(
            content_id=outside_content.id,
            q_type="single",
            question=f"{prefix}-outside-question",
            options=["A", "B"],
            answer="A",
            created_at=datetime(2026, 1, 2, tzinfo=timezone.utc),
        ))
        await db.commit()
        collection_id = str(collection.id)

    page_1 = await client.get(f"/api/ai/quiz/history?scope_type=collection&scope_id={collection_id}&page=1&page_size=3")
    page_3 = await client.get(f"/api/ai/quiz/history?scope_type=collection&scope_id={collection_id}&page=3&page_size=3")

    assert page_1.status_code == 200
    page_1_body = page_1.json()
    assert page_1_body["total"] == 7
    assert page_1_body["page"] == 1
    assert page_1_body["page_size"] == 3
    assert len(page_1_body["questions"]) == 3
    assert all(prefix in item["question"] and "outside" not in item["question"] for item in page_1_body["questions"])

    assert page_3.status_code == 200
    page_3_body = page_3.json()
    assert page_3_body["total"] == 7
    assert page_3_body["page"] == 3
    assert len(page_3_body["questions"]) == 1
    assert "outside" not in page_3_body["questions"][0]["question"]


@pytest.mark.asyncio
async def test_ai_manual_scope_rejects_cross_brain_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content

    brain_a_resp = await client.post("/api/brains", json={"name": f"ai-scope-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"ai-scope-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = uuid.UUID(brain_b_resp.json()["id"])

    async with async_session_factory() as db:
        content = Content(
            title=f"cross-brain-ai-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="brain b content",
            processing_status="completed",
            brain_id=brain_b,
        )
        db.add(content)
        await db.commit()
        content_id = str(content.id)

    resp = await client.post(
        "/api/ai/quiz",
        json={"content_ids": [content_id], "brain_id": brain_a},
    )
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_ai_ask_empty_brain_returns_no_content_without_embedding(client: AsyncClient):
    brain_resp = await client.post("/api/brains", json={"name": f"empty-qa-{uuid.uuid4().hex[:8]}"})
    assert brain_resp.status_code == 200
    brain_id = brain_resp.json()["id"]

    resp = await client.post(
        "/api/ai/ask",
        json={"question": "这个工作区有什么？", "brain_id": brain_id},
    )
    assert resp.status_code == 200
    assert resp.json()["answer"] == "所选范围内没有可用的内容"
    assert resp.json()["sources"] == []


@pytest.mark.asyncio
async def test_prompt_templates_are_brain_scoped(client: AsyncClient):
    brain_a_resp = await client.post("/api/brains", json={"name": f"brain-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"brain-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]
    prompt_a = f"system-a-{uuid.uuid4().hex[:8]}"
    prompt_b = f"system-b-{uuid.uuid4().hex[:8]}"

    resp_a = await client.put(
        "/api/ai/quiz-template",
        json={"system_prompt": prompt_a, "user_prompt_template": "user-a", "brain_id": brain_a},
    )
    resp_b = await client.put(
        "/api/ai/quiz-template",
        json={"system_prompt": prompt_b, "user_prompt_template": "user-b", "brain_id": brain_b},
    )
    assert resp_a.status_code == 200
    assert resp_b.status_code == 200

    get_a = await client.get(f"/api/ai/quiz-template?brain_id={brain_a}")
    get_b = await client.get(f"/api/ai/quiz-template?brain_id={brain_b}")
    assert get_a.status_code == 200
    assert get_b.status_code == 200
    assert get_a.json()["template"]["system_prompt"] == prompt_a
    assert get_b.json()["template"]["system_prompt"] == prompt_b


@pytest.mark.asyncio
async def test_brain_config_overrides_ai_provider(client: AsyncClient):
    from app.api.ai import _get_ai_provider
    from app.core.database import async_session_factory
    from app.models.models import Brain, ProviderConfig

    async with async_session_factory() as db:
        provider = ProviderConfig(
            name=f"pytest-brain-ai-provider-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            base_url="https://example.test/v1",
            default_models={"summarize": "provider-summary"},
            is_active=True,
        )
        brain = Brain(
            name=f"pytest-brain-ai-config-{uuid.uuid4().hex[:8]}",
            config={"provider_id": "", "summarize_model": "brain-summary", "qa_model": "brain-qa", "judge_model": "brain-judge"},
        )
        db.add_all([provider, brain])
        await db.flush()
        brain.config = {
            "provider_id": str(provider.id),
            "summarize_model": "brain-summary",
            "qa_model": "brain-qa",
            "judge_model": "brain-judge",
        }
        provider_config = await _get_ai_provider(db, "summarize", brain.id)
        qa_provider_config = await _get_ai_provider(db, "qa", brain.id)
        judge_provider_config = await _get_ai_provider(db, "judge", brain.id)

    assert provider_config is not None
    assert provider_config["provider_id"] == str(provider.id)
    assert provider_config["model"] == "brain-summary"
    assert qa_provider_config is not None
    assert qa_provider_config["provider_id"] == str(provider.id)
    assert qa_provider_config["model"] == "brain-qa"
    assert judge_provider_config is not None
    assert judge_provider_config["provider_id"] == str(provider.id)
    assert judge_provider_config["model"] == "brain-judge"


@pytest.mark.asyncio
async def test_brain_config_overrides_embedding_binding(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, ProviderConfig
    from app.services.embedding import _get_embedding_binding

    async with async_session_factory() as db:
        provider = ProviderConfig(
            name=f"pytest-brain-embed-provider-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            base_url="https://example.test/v1",
            default_models={"embedding": "provider-embedding"},
            is_active=True,
        )
        brain = Brain(name=f"pytest-brain-embed-config-{uuid.uuid4().hex[:8]}")
        db.add_all([provider, brain])
        await db.flush()
        brain.config = {"provider_id": str(provider.id), "embedding_model": "brain-embedding"}
        binding = await _get_embedding_binding(db, brain.id)

    assert binding == {"provider_id": str(provider.id), "model": "brain-embedding"}


@pytest.mark.asyncio
async def test_brain_config_validates_provider_and_cleans_empty_values(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, ProviderConfig

    brain_resp = await client.post("/api/brains", json={"name": f"config-brain-{uuid.uuid4().hex[:8]}"})
    assert brain_resp.status_code == 200
    brain_id = brain_resp.json()["id"]

    missing_resp = await client.put(
        f"/api/brains/{brain_id}/config",
        json={"provider_id": str(uuid.uuid4()), "summarize_model": "gpt-test"},
    )
    assert missing_resp.status_code == 404

    empty_resp = await client.put(
        f"/api/brains/{brain_id}/config",
        json={"provider_id": "", "summarize_model": " "},
    )
    assert empty_resp.status_code == 200
    empty_config = await client.get(f"/api/brains/{brain_id}/config")
    assert empty_config.status_code == 200
    assert empty_config.json() == {}

    async with async_session_factory() as db:
        provider = ProviderConfig(
            name=f"pytest-brain-config-provider-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            base_url="https://example.test/v1",
            is_active=True,
        )
        db.add(provider)
        await db.commit()
        provider_id = str(provider.id)

    valid_resp = await client.put(
        f"/api/brains/{brain_id}/config",
        json={
            "provider_id": provider_id,
            "summarize_model": "  gpt-trimmed  ",
            "qa_model": "  gpt-qa  ",
            "judge_model": "  gpt-judge  ",
        },
    )
    assert valid_resp.status_code == 200
    valid_config = await client.get(f"/api/brains/{brain_id}/config")
    assert valid_config.json() == {
        "provider_id": provider_id,
        "summarize_model": "gpt-trimmed",
        "qa_model": "gpt-qa",
        "judge_model": "gpt-judge",
    }


@pytest.mark.asyncio
async def test_study_brain_template_creates_structure(client: AsyncClient):
    resp = await client.post(
        "/api/brains",
        json={"name": f"study-brain-{uuid.uuid4().hex[:8]}", "template": "study"},
    )
    assert resp.status_code == 200
    brain_id = resp.json()["id"]

    cats = await client.get(f"/api/categories?brain_id={brain_id}")
    cols = await client.get(f"/api/collections?brain_id={brain_id}")
    assert cats.status_code == 200
    assert cols.status_code == 200

    category_names = {item["name"] for item in cats.json()}
    collection_names = {item["name"] for item in cols.json()["items"]}
    assert {"未分类", "基础概念", "课程笔记", "资料阅读", "实践练习", "复盘总结"}.issubset(category_names)
    assert {"课程合集", "书籍与 PDF", "案例与练习"}.issubset(collection_names)

    tags = await client.get(f"/api/tags?brain_id={brain_id}")
    assert tags.status_code == 200
    tag_names = {item["name"] for item in tags.json()}
    assert {"入门", "重点", "待复习", "已掌握", "需要实践", "作业"}.issubset(tag_names)


@pytest.mark.asyncio
async def test_tags_are_unique_within_brain_not_global(client: AsyncClient):
    tag_name = f"shared-tag-{uuid.uuid4().hex[:8]}"
    brain_a_resp = await client.post("/api/brains", json={"name": f"tag-brain-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"tag-brain-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]

    first = await client.post("/api/tags", json={"name": tag_name, "brain_id": brain_a})
    same_brain = await client.post("/api/tags", json={"name": tag_name, "brain_id": brain_a})
    other_brain = await client.post("/api/tags", json={"name": tag_name, "brain_id": brain_b})

    assert first.status_code == 201
    assert same_brain.status_code == 409
    assert other_brain.status_code == 201


@pytest.mark.asyncio
async def test_brain_overview_returns_stats_and_recent_contents(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Category, Collection, CollectionItem, Content, Tag

    create_resp = await client.post(
        "/api/brains",
        json={"name": f"overview-brain-{uuid.uuid4().hex[:8]}", "template": "blank"},
    )
    assert create_resp.status_code == 200
    brain_id = uuid.UUID(create_resp.json()["id"])

    completed_title = f"overview-completed-{uuid.uuid4().hex[:8]}"
    failed_title = f"overview-failed-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        completed_content = Content(
            title=completed_title,
            content_type="pdf",
            source_type="upload",
            processing_status="completed",
            file_size=120,
            brain_id=brain_id,
            extra_meta={"study_status": "completed"},
        )
        resume_content = Content(
            title=failed_title,
            content_type="video",
            source_type="upload",
            processing_status="failed",
            file_size=80,
            brain_id=brain_id,
            extra_meta={"study_status": "in_progress"},
        )
        resume_collection = Collection(name=f"overview-col-{uuid.uuid4().hex[:8]}", brain_id=brain_id)
        db.add_all([
            completed_content,
            resume_content,
            Category(name=f"overview-cat-{uuid.uuid4().hex[:8]}", brain_id=brain_id),
            Tag(name=f"overview-tag-{uuid.uuid4().hex[:8]}", brain_id=brain_id),
            resume_collection,
        ])
        await db.flush()
        db.add(CollectionItem(collection_id=resume_collection.id, content_id=resume_content.id, sort_order=1))
        await db.commit()
        resume_collection_id = str(resume_collection.id)
        resume_collection_name = resume_collection.name

    overview_resp = await client.get(f"/api/brains/{brain_id}/overview")
    assert overview_resp.status_code == 200
    data = overview_resp.json()
    stats = data["stats"]

    assert data["brain"]["id"] == str(brain_id)
    assert stats["total_contents"] == 2
    assert stats["storage_bytes"] == 200
    assert stats["by_status"]["completed"] == 1
    assert stats["by_status"]["failed"] == 1
    assert stats["by_type"]["pdf"] == 1
    assert stats["by_type"]["video"] == 1
    assert stats["categories"] >= 2
    assert stats["tags"] == 1
    assert stats["collections"] == 1
    assert data["study"] == {
        "total": 2,
        "completed": 1,
        "in_progress": 1,
        "not_started": 0,
        "progress_percent": 50,
    }
    assert data["resume_content"]["title"] == failed_title
    assert data["resume_content"]["study_status"] == "in_progress"
    assert data["resume_content"]["collection_id"] == resume_collection_id
    assert data["resume_content"]["collection_name"] == resume_collection_name
    assert {item["title"] for item in data["recent_contents"]} >= {completed_title, failed_title}


@pytest.mark.asyncio
async def test_content_brain_move_cleans_old_organization_links(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import (
        Category,
        Collection,
        CollectionItem,
        Content,
        ContentCategory,
        ContentRelation,
        ContentTag,
        Tag,
    )

    brain_a_resp = await client.post("/api/brains", json={"name": f"move-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"move-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = uuid.UUID(brain_a_resp.json()["id"])
    brain_b = uuid.UUID(brain_b_resp.json()["id"])

    async with async_session_factory() as db:
        content = Content(
            title=f"move-content-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="move me",
            brain_id=brain_a,
        )
        related = Content(
            title=f"move-related-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="related",
            brain_id=brain_a,
        )
        category = Category(name=f"move-cat-{uuid.uuid4().hex[:8]}", brain_id=brain_a)
        tag = Tag(name=f"move-tag-{uuid.uuid4().hex[:8]}", brain_id=brain_a)
        collection = Collection(name=f"move-col-{uuid.uuid4().hex[:8]}", brain_id=brain_a)
        db.add_all([content, related, category, tag, collection])
        await db.flush()
        db.add_all([
            ContentCategory(content_id=content.id, category_id=category.id),
            ContentTag(content_id=content.id, tag_id=tag.id),
            CollectionItem(content_id=content.id, collection_id=collection.id),
            ContentRelation(source_id=content.id, target_id=related.id, relation_type="reference"),
        ])
        await db.commit()
        content_id = str(content.id)

    move_resp = await client.patch(f"/api/contents/{content_id}", json={"brain_id": str(brain_b)})
    assert move_resp.status_code == 200
    assert move_resp.json()["brain_id"] == str(brain_b)

    async with async_session_factory() as db:
        moved = await db.get(Content, uuid.UUID(content_id))
        category_links = await db.execute(select(ContentCategory).where(ContentCategory.content_id == moved.id))
        tag_links = await db.execute(select(ContentTag).where(ContentTag.content_id == moved.id))
        collection_links = await db.execute(select(CollectionItem).where(CollectionItem.content_id == moved.id))
        relation_links = await db.execute(
            select(ContentRelation).where(
                (ContentRelation.source_id == moved.id) | (ContentRelation.target_id == moved.id)
            )
        )

    assert moved.brain_id == brain_b
    assert category_links.scalars().all() == []
    assert tag_links.scalars().all() == []
    assert collection_links.scalars().all() == []
    assert relation_links.scalars().all() == []


@pytest.mark.asyncio
async def test_batch_content_move_scopes_source_and_cleans_old_links(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Category, Collection, CollectionItem, Content, ContentCategory

    brain_a_resp = await client.post("/api/brains", json={"name": f"batch-move-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"batch-move-b-{uuid.uuid4().hex[:8]}"})
    brain_c_resp = await client.post("/api/brains", json={"name": f"batch-move-c-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    assert brain_c_resp.status_code == 200
    brain_a = uuid.UUID(brain_a_resp.json()["id"])
    brain_b = uuid.UUID(brain_b_resp.json()["id"])
    brain_c = uuid.UUID(brain_c_resp.json()["id"])

    async with async_session_factory() as db:
        content_a = Content(
            title=f"batch-move-content-a-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="move me",
            brain_id=brain_a,
        )
        content_c = Content(
            title=f"batch-move-content-c-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="do not move me",
            brain_id=brain_c,
        )
        category = Category(name=f"batch-move-cat-{uuid.uuid4().hex[:8]}", brain_id=brain_a)
        collection = Collection(name=f"batch-move-col-{uuid.uuid4().hex[:8]}", brain_id=brain_a)
        db.add_all([content_a, content_c, category, collection])
        await db.flush()
        db.add_all([
            ContentCategory(content_id=content_a.id, category_id=category.id),
            CollectionItem(content_id=content_a.id, collection_id=collection.id),
        ])
        await db.commit()
        content_a_id = content_a.id
        content_c_id = content_c.id

    move_resp = await client.post(
        "/api/contents/batch-move",
        json={
            "ids": [str(content_a_id), str(content_c_id)],
            "brain_id": str(brain_a),
            "target_brain_id": str(brain_b),
        },
    )
    missing_target_resp = await client.post(
        "/api/contents/batch-move",
        json={"ids": [], "target_brain_id": str(uuid.uuid4())},
    )

    assert move_resp.status_code == 200
    assert move_resp.json()["moved"] == 1
    assert missing_target_resp.status_code == 404

    async with async_session_factory() as db:
        moved = await db.get(Content, content_a_id)
        skipped = await db.get(Content, content_c_id)
        category_links = await db.execute(select(ContentCategory).where(ContentCategory.content_id == content_a_id))
        collection_links = await db.execute(select(CollectionItem).where(CollectionItem.content_id == content_a_id))

    assert moved.brain_id == brain_b
    assert skipped.brain_id == brain_c
    assert category_links.scalars().all() == []
    assert collection_links.scalars().all() == []


@pytest.mark.asyncio
async def test_content_write_rejects_unknown_brain_id(client: AsyncClient):
    missing_brain = uuid.uuid4()

    create_resp = await client.post(
        "/api/contents",
        json={
            "title": f"unknown-brain-create-{uuid.uuid4().hex[:8]}",
            "content_type": "note",
            "text_content": "orphan guard",
            "brain_id": str(missing_brain),
        },
    )
    assert create_resp.status_code == 404

    upload_resp = await client.post(
        "/api/files/upload",
        data={"brain_id": str(missing_brain)},
        files={"file": ("unknown-brain.txt", b"orphan upload guard", "text/plain")},
    )
    assert upload_resp.status_code == 404

    valid_resp = await client.post(
        "/api/contents",
        json={
            "title": f"unknown-brain-update-{uuid.uuid4().hex[:8]}",
            "content_type": "note",
            "text_content": "move guard",
        },
    )
    assert valid_resp.status_code == 201

    update_resp = await client.patch(
        f"/api/contents/{valid_resp.json()['id']}",
        json={"brain_id": str(missing_brain)},
    )
    assert update_resp.status_code == 404


@pytest.mark.asyncio
async def test_content_update_merges_extra_meta(client: AsyncClient):
    original_meta = {
        "import_relative_path": "Course/Week 1/lesson.txt",
        "import_root": "Course",
    }
    create_resp = await client.post(
        "/api/contents",
        json={
            "title": f"meta-merge-{uuid.uuid4().hex[:8]}",
            "content_type": "doc",
            "source_type": "manual",
            "extra_meta": original_meta,
        },
    )
    assert create_resp.status_code == 201
    content_id = create_resp.json()["id"]

    update_resp = await client.patch(
        f"/api/contents/{content_id}",
        json={"extra_meta": {"study_status": "completed"}},
    )

    assert update_resp.status_code == 200
    meta = update_resp.json()["extra_meta"]
    assert meta["import_relative_path"] == original_meta["import_relative_path"]
    assert meta["import_root"] == original_meta["import_root"]
    assert meta["study_status"] == "completed"


@pytest.mark.asyncio
async def test_batch_study_status_merges_extra_meta_and_scopes_brain(client: AsyncClient):
    brain_a_resp = await client.post("/api/brains", json={"name": f"study-batch-a-{uuid.uuid4().hex[:8]}"})
    brain_b_resp = await client.post("/api/brains", json={"name": f"study-batch-b-{uuid.uuid4().hex[:8]}"})
    assert brain_a_resp.status_code == 200
    assert brain_b_resp.status_code == 200
    brain_a = brain_a_resp.json()["id"]
    brain_b = brain_b_resp.json()["id"]

    create_a_resp = await client.post(
        "/api/contents",
        json={
            "title": f"batch-study-a-{uuid.uuid4().hex[:8]}",
            "content_type": "doc",
            "source_type": "manual",
            "brain_id": brain_a,
            "extra_meta": {
                "import_relative_path": "Course/Week 2/lesson 10.txt",
                "study_started_at": "2026-01-01T00:00:00+00:00",
            },
        },
    )
    create_b_resp = await client.post(
        "/api/contents",
        json={
            "title": f"batch-study-b-{uuid.uuid4().hex[:8]}",
            "content_type": "doc",
            "source_type": "manual",
            "brain_id": brain_b,
            "extra_meta": {"import_relative_path": "Other/lesson.txt"},
        },
    )
    assert create_a_resp.status_code == 201
    assert create_b_resp.status_code == 201
    content_a = create_a_resp.json()["id"]
    content_b = create_b_resp.json()["id"]

    batch_resp = await client.post(
        "/api/contents/batch-study-status",
        json={"ids": [content_a, content_b], "status": "completed", "brain_id": brain_a},
    )
    invalid_status_resp = await client.post(
        "/api/contents/batch-study-status",
        json={"ids": [content_a], "status": "paused", "brain_id": brain_a},
    )

    assert batch_resp.status_code == 200
    assert batch_resp.json()["updated"] == 1
    assert invalid_status_resp.status_code == 400

    detail_a = await client.get(f"/api/files/{content_a}")
    detail_b = await client.get(f"/api/files/{content_b}")
    meta_a = detail_a.json()["extra_meta"]
    meta_b = detail_b.json()["extra_meta"]
    assert meta_a["import_relative_path"] == "Course/Week 2/lesson 10.txt"
    assert meta_a["study_status"] == "completed"
    assert meta_a["study_started_at"] == "2026-01-01T00:00:00+00:00"
    assert meta_a["study_completed_at"] is not None
    assert "study_status" not in meta_b


@pytest.mark.asyncio
async def test_content_reads_and_batch_actions_reject_unknown_brain_id(client: AsyncClient):
    missing_brain = uuid.uuid4()

    files_resp = await client.get(f"/api/files?brain_id={missing_brain}")
    storage_resp = await client.get(f"/api/storage/stats?brain_id={missing_brain}")
    batch_chunk_resp = await client.post(f"/api/contents/batch-chunk?brain_id={missing_brain}", json=[])
    batch_embed_resp = await client.post(f"/api/contents/batch-embed?brain_id={missing_brain}", json=[])
    rechunk_resp = await client.post(f"/api/contents/rechunk-all?brain_id={missing_brain}")
    batch_action_resp = await client.post(
        "/api/contents/batch",
        json={"ids": [], "action": "star", "brain_id": str(missing_brain)},
    )
    batch_study_resp = await client.post(
        "/api/contents/batch-study-status",
        json={"ids": [], "status": "completed", "brain_id": str(missing_brain)},
    )
    reset_resp = await client.post(f"/api/contents/maintenance/reset-stuck-embeddings?brain_id={missing_brain}")

    assert files_resp.status_code == 404
    assert storage_resp.status_code == 404
    assert batch_chunk_resp.status_code == 404
    assert batch_embed_resp.status_code == 404
    assert rechunk_resp.status_code == 404
    assert batch_action_resp.status_code == 404
    assert batch_study_resp.status_code == 404
    assert reset_resp.status_code == 404


@pytest.mark.asyncio
async def test_delete_brain_removes_related_content_and_organization(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import (
        Brain,
        Category,
        Collection,
        CollectionItem,
        Content,
        ContentCategory,
        ContentTag,
        PromptTemplate,
        SearchLog,
        Tag,
    )
    from sqlalchemy import func

    settings = get_settings()
    old_storage_root = settings.file_storage_root
    storage_root = tmp_path / "storage"
    storage_root.mkdir()
    file_path = storage_root / "delete-brain" / "stored.txt"
    file_path.parent.mkdir()
    file_path.write_text("delete me", encoding="utf-8")
    try:
        settings.file_storage_root = str(storage_root)
        create_resp = await client.post("/api/brains", json={"name": f"delete-brain-{uuid.uuid4().hex[:8]}"})
        assert create_resp.status_code == 200
        brain_id = uuid.UUID(create_resp.json()["id"])

        async with async_session_factory() as db:
            content = Content(
                title=f"delete-content-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="upload",
                file_path="delete-brain/stored.txt",
                brain_id=brain_id,
            )
            category = Category(name=f"delete-cat-{uuid.uuid4().hex[:8]}", brain_id=brain_id)
            tag = Tag(name=f"delete-tag-{uuid.uuid4().hex[:8]}", brain_id=brain_id)
            collection = Collection(name=f"delete-col-{uuid.uuid4().hex[:8]}", brain_id=brain_id)
            db.add_all([
                content,
                category,
                tag,
                collection,
                SearchLog(query=f"delete-log-{uuid.uuid4().hex[:8]}", brain_id=brain_id),
            ])
            await db.flush()
            db.add_all([
                ContentCategory(content_id=content.id, category_id=category.id),
                ContentTag(content_id=content.id, tag_id=tag.id),
                CollectionItem(content_id=content.id, collection_id=collection.id),
            ])
            await db.commit()

        delete_resp = await client.delete(f"/api/brains/{brain_id}")
        assert delete_resp.status_code == 200
        assert delete_resp.json()["deleted_contents"] == 1
        assert delete_resp.json()["removed_files"] == 1
        assert not file_path.exists()

        async with async_session_factory() as db:
            for model in (Brain, Content, Category, Tag, Collection, PromptTemplate, SearchLog):
                result = await db.execute(select(func.count()).select_from(model).where(model.brain_id == brain_id) if model is not Brain else select(func.count()).select_from(model).where(model.id == brain_id))
                assert result.scalar() == 0
    finally:
        settings.file_storage_root = old_storage_root


# ── Search ──

@pytest.mark.asyncio
async def test_search_keyword(client: AsyncClient):
    resp = await client.post(
        "/api/search",
        json={"query": "test", "top_k": 5, "enable_vector": False},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert "results" in data


# ── Notes ──

@pytest.mark.asyncio
async def test_search_filters_by_created_date_range(client: AsyncClient):
    from datetime import datetime, timedelta, timezone
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentChunk

    old_title = f"pytest-old-search-{uuid.uuid4().hex[:8]}"
    new_title = f"pytest-new-search-{uuid.uuid4().hex[:8]}"
    token = f"daterange-{uuid.uuid4().hex[:8]}"
    now = datetime.now(timezone.utc)
    old_date = now - timedelta(days=60)

    async with async_session_factory() as db:
        old_content = Content(
            title=old_title,
            content_type="doc",
            source_type="manual",
            text_content=token,
            processing_status="completed",
            created_at=old_date,
        )
        new_content = Content(
            title=new_title,
            content_type="doc",
            source_type="manual",
            text_content=token,
            processing_status="completed",
            created_at=now,
        )
        db.add_all([old_content, new_content])
        await db.flush()
        db.add_all([
            ContentChunk(content_id=old_content.id, chunk_index=0, chunk_type="text", chunk_text=token),
            ContentChunk(content_id=new_content.id, chunk_index=0, chunk_type="text", chunk_text=token),
        ])
        await db.commit()

    resp = await client.post(
        "/api/search",
        json={
            "query": token,
            "top_k": 10,
            "enable_vector": False,
            "created_after": (now - timedelta(days=7)).isoformat(),
        },
    )
    assert resp.status_code == 200
    titles = {item["title"] for item in resp.json()["results"]}
    assert new_title in titles
    assert old_title not in titles


@pytest.mark.asyncio
async def test_search_filters_by_tag_category_and_brain(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, Category, Content, ContentCategory, ContentChunk, ContentTag, Tag

    token = f"scoped-filter-{uuid.uuid4().hex[:8]}"
    target_title = f"pytest-filter-target-{uuid.uuid4().hex[:8]}"
    other_title = f"pytest-filter-other-{uuid.uuid4().hex[:8]}"
    brain_id = uuid.uuid4()

    async with async_session_factory() as db:
        brain = Brain(id=brain_id, name=f"search-filter-brain-{uuid.uuid4().hex[:8]}")
        tag = Tag(name=f"pytest-tag-filter-{uuid.uuid4().hex[:8]}")
        category = Category(name=f"pytest-category-filter-{uuid.uuid4().hex[:8]}")
        target = Content(
            title=target_title,
            content_type="doc",
            source_type="manual",
            text_content=token,
            processing_status="completed",
            brain_id=brain_id,
        )
        other = Content(
            title=other_title,
            content_type="doc",
            source_type="manual",
            text_content=token,
            processing_status="completed",
        )
        db.add_all([brain, tag, category, target, other])
        await db.flush()
        db.add_all([
            ContentTag(content_id=target.id, tag_id=tag.id),
            ContentCategory(content_id=target.id, category_id=category.id),
            ContentChunk(content_id=target.id, chunk_index=0, chunk_type="text", chunk_text=token),
            ContentChunk(content_id=other.id, chunk_index=0, chunk_type="text", chunk_text=token),
        ])
        await db.commit()
        tag_id = str(tag.id)
        category_id = str(category.id)

    base = {"query": token, "top_k": 10, "enable_vector": False}
    cases = [
        {"tag_ids": [tag_id]},
        {"category_id": category_id},
        {"brain_id": str(brain_id)},
    ]
    for filters in cases:
        resp = await client.post("/api/search", json={**base, **filters})
        assert resp.status_code == 200
        titles = {item["title"] for item in resp.json()["results"]}
        assert target_title in titles
        assert other_title not in titles


@pytest.mark.asyncio
async def test_note_crud(client: AsyncClient):
    resp = await client.post("/api/notes", json={"title": "TestNote", "content": "hello"})
    assert resp.status_code == 200
    nid = resp.json()["id"]

    resp2 = await client.get(f"/api/notes/{nid}")
    assert resp2.status_code == 200
    assert resp2.json()["title"] == "TestNote"

    resp3 = await client.delete(f"/api/notes/{nid}")
    assert resp3.status_code == 200


# ── Embedding ──

@pytest.mark.asyncio
async def test_web_content_create_and_chunk(client: AsyncClient):
    text = "Moyuan web capture test content. " * 12
    resp = await client.post(
        "/api/contents",
        json={
            "title": f"pytest-web-{uuid.uuid4().hex[:8]}",
            "content_type": "web",
            "source_type": "web_capture",
            "source_url": "https://example.com/moyuan-test",
            "text_content": text,
        },
    )
    assert resp.status_code == 201
    content_id = resp.json()["id"]

    process_resp = await client.post(f"/api/contents/{content_id}/process")
    assert process_resp.status_code == 200
    assert process_resp.json()["processing_status"] in {"completed", "chunked"}

    chunks_resp = await client.get(f"/api/contents/{content_id}/chunks")
    assert chunks_resp.status_code == 200
    data = chunks_resp.json()
    assert data["total"] >= 1
    assert data["chunks"][0]["chunk_text"]


@pytest.mark.asyncio
async def test_web_preview_fetches_text_without_creating_content(client: AsyncClient, monkeypatch):
    import app.services.process as process_module

    async def fake_extract_web(url):
        assert url == "https://example.com/article"
        return "Preview Article Title\n\nMoyuan preview body text."

    monkeypatch.setattr(process_module, "_extract_web", fake_extract_web)

    resp = await client.post("/api/contents/web-preview", json={"url": "https://example.com/article"})
    assert resp.status_code == 200
    data = resp.json()
    assert data["title"] == "Preview Article Title"
    assert "Moyuan preview body text" in data["text_content"]
    assert data["text_length"] > 0


@pytest.mark.asyncio
async def test_web_preview_rejects_invalid_url(client: AsyncClient):
    resp = await client.post("/api/contents/web-preview", json={"url": "not-a-url"})
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_web_content_adds_screenshot_chunk(client: AsyncClient, monkeypatch):
    import app.services.process as process_module

    async def fake_capture(url, output_dir):
        assert url == "https://example.com/moyuan-screenshot"
        assert output_dir.name
        return "images/web_capture_test.png"

    monkeypatch.setattr(process_module, "_capture_web_screenshot", fake_capture)

    text = "Moyuan web capture screenshot content. " * 12
    resp = await client.post(
        "/api/contents",
        json={
            "title": f"pytest-web-shot-{uuid.uuid4().hex[:8]}",
            "content_type": "web",
            "source_type": "web_capture",
            "source_url": "https://example.com/moyuan-screenshot",
            "text_content": text,
        },
    )
    assert resp.status_code == 201
    content_id = resp.json()["id"]

    process_resp = await client.post(f"/api/contents/{content_id}/process")
    assert process_resp.status_code == 200

    chunks_resp = await client.get(f"/api/contents/{content_id}/chunks")
    assert chunks_resp.status_code == 200
    chunks = chunks_resp.json()["chunks"]
    assert any(chunk["chunk_type"] == "text" and chunk["chunk_text"] for chunk in chunks)
    assert chunks[-1]["chunk_type"] == "image"
    assert chunks[-1]["image_path"] == "images/web_capture_test.png"


@pytest.mark.asyncio
async def test_image_processing_adds_ocr_text_chunk(client: AsyncClient, monkeypatch, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    import app.services.process as process_module

    settings = get_settings()
    image_dir = tmp_path / "uploads"
    image_dir.mkdir()
    image_file = image_dir / "ocr-test.png"
    image_file.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00"
        b"\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00"
        b"\x05\xfe\x02\xfeA\xe2!\xbc\x00\x00\x00\x00IEND\xaeB`\x82"
    )

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(tmp_path)

    async def fake_ocr(path, db=None):
        assert path.name == "ocr-test.png"
        return "Moyuan OCR extracted text"

    monkeypatch.setattr(process_module, "_ocr_image", fake_ocr)

    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-image-{uuid.uuid4().hex[:8]}",
                content_type="image",
                source_type="upload",
                file_path="uploads/ocr-test.png",
            )
            db.add(content)
            await db.commit()
            content_id = str(content.id)

        process_resp = await client.post(f"/api/contents/{content_id}/process")
        assert process_resp.status_code == 200
        assert process_resp.json()["processing_status"] in {"completed", "chunked"}

        chunks_resp = await client.get(f"/api/contents/{content_id}/chunks")
        assert chunks_resp.status_code == 200
        chunks = chunks_resp.json()["chunks"]
        assert [chunk["chunk_type"] for chunk in chunks] == ["image", "text"]
        assert chunks[1]["chunk_text"] == "Moyuan OCR extracted text"

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == uuid.UUID(content_id)))
            refreshed = result.scalar_one()
            assert refreshed.text_content == "Moyuan OCR extracted text"
    finally:
        settings.file_storage_root = old_storage_root


@pytest.mark.asyncio
async def test_ocr_provider_uses_configured_openai_compatible_api(client: AsyncClient, monkeypatch, tmp_path):
    from app.core.database import async_session_factory
    from app.models.models import FunctionBindingConfig, ProviderConfig
    from app.core.crypto import crypto_service
    from sqlalchemy import delete
    import app.services.process as process_module

    image_path = tmp_path / "configured-ocr.png"
    image_path.write_bytes(b"fake-image")
    captured: dict[str, str] = {}

    class FakeMessage:
        content = "Configured OCR line"

    class FakeChoice:
        message = FakeMessage()

    class FakeResponse:
        choices = [FakeChoice()]

    class FakeCompletions:
        async def create(self, **kwargs):
            captured["model"] = kwargs["model"]
            captured["messages"] = str(kwargs["messages"])
            return FakeResponse()

    class FakeChat:
        completions = FakeCompletions()

    class FakeOpenAIClient:
        chat = FakeChat()

        def __init__(self, api_key, base_url):
            captured["api_key"] = api_key
            captured["base_url"] = base_url

    import openai
    monkeypatch.setattr(openai, "AsyncOpenAI", FakeOpenAIClient)

    async with async_session_factory() as db:
        await db.execute(delete(FunctionBindingConfig).where(FunctionBindingConfig.function == "ocr"))
        provider = ProviderConfig(
            name=f"pytest-configured-ocr-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            api_key_encrypted=crypto_service.encrypt("secret-key"),
            base_url="https://example.test/v1",
            is_active=True,
        )
        db.add(provider)
        await db.flush()
        db.add(FunctionBindingConfig(function="ocr", provider_id=provider.id, model="gpt-vision-test"))
        await db.commit()

        text = await process_module._ocr_image(image_path, db)

    assert text == "Configured OCR line"
    assert captured["api_key"] == "secret-key"
    assert captured["base_url"] == "https://example.test/v1"
    assert captured["model"] == "gpt-vision-test"
    assert captured["base_url"] == "https://example.test/v1"
    assert "data:image" in captured["messages"]


@pytest.mark.asyncio
async def test_audio_processing_adds_transcript_chunks(client: AsyncClient, monkeypatch, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    import app.services.process as process_module

    settings = get_settings()
    audio_dir = tmp_path / "uploads"
    audio_dir.mkdir()
    audio_file = audio_dir / "transcript-test.mp3"
    audio_file.write_bytes(b"fake-audio-for-transcript-test")

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(tmp_path)

    async def fake_transcribe(path, db=None):
        assert path.name == "transcript-test.mp3"
        return [
            {"text": "First audio segment", "start": 0.0, "end": 1.5},
            {"text": "Second audio segment", "start": 1.5, "end": 3.0},
        ]

    monkeypatch.setattr(process_module, "_transcribe_audio", fake_transcribe)

    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-audio-{uuid.uuid4().hex[:8]}",
                content_type="audio",
                source_type="upload",
                file_path="uploads/transcript-test.mp3",
            )
            db.add(content)
            await db.commit()
            content_id = str(content.id)

        process_resp = await client.post(f"/api/contents/{content_id}/process")
        assert process_resp.status_code == 200
        assert process_resp.json()["processing_status"] in {"completed", "chunked"}

        chunks_resp = await client.get(f"/api/contents/{content_id}/chunks")
        assert chunks_resp.status_code == 200
        chunks = chunks_resp.json()["chunks"]
        assert [chunk["chunk_text"] for chunk in chunks] == [
            "First audio segment",
            "Second audio segment",
        ]
        assert chunks[0]["time_start"] == 0.0
        assert chunks[1]["time_end"] == 3.0

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == uuid.UUID(content_id)))
            refreshed = result.scalar_one()
            assert refreshed.text_content == "First audio segment Second audio segment"
    finally:
        settings.file_storage_root = old_storage_root


@pytest.mark.asyncio
async def test_local_faster_whisper_transcription_binding(client: AsyncClient, monkeypatch, tmp_path):
    from app.core.database import async_session_factory
    from app.models.models import FunctionBindingConfig, ProviderConfig
    from sqlalchemy import delete
    import app.services.process as process_module

    audio_path = tmp_path / "local-whisper.mp3"
    audio_path.write_bytes(b"fake-local-whisper-audio")

    async def fake_local_transcribe(path, model, extra_params=None):
        assert path == audio_path
        assert model == "tiny"
        assert extra_params["transcribe_backend"] == "faster_whisper"
        assert extra_params["language"] == "zh"
        return [{"text": "Local faster whisper text", "start": 0.5, "end": 2.0}]

    monkeypatch.setattr(process_module, "_transcribe_with_faster_whisper", fake_local_transcribe)

    provider_name = f"pytest-local-whisper-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        await db.execute(delete(FunctionBindingConfig).where(FunctionBindingConfig.function == "transcribe"))
        provider = ProviderConfig(
            name=provider_name,
            provider_type="custom",
            base_url="local:faster-whisper",
            extra_params={"transcribe_backend": "faster_whisper", "language": "zh"},
            is_active=True,
        )
        db.add(provider)
        await db.flush()
        db.add(FunctionBindingConfig(function="transcribe", provider_id=provider.id, model="tiny"))
        await db.commit()

        segments = await process_module._transcribe_audio(audio_path, db)

    assert segments == [{"text": "Local faster whisper text", "start": 0.5, "end": 2.0}]


@pytest.mark.asyncio
async def test_video_processing_adds_transcript_and_screenshot_chunks(client: AsyncClient, monkeypatch, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    import app.services.process as process_module

    settings = get_settings()
    video_dir = tmp_path / "uploads"
    video_dir.mkdir()
    video_file = video_dir / "video-test.mp4"
    video_file.write_bytes(b"fake-video-for-process-test")

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(tmp_path)

    async def fake_process_video(path, db=None, screenshot_dir=None):
        assert path.name == "video-test.mp4"
        assert screenshot_dir is not None
        return (
            [
                {"text": "First video segment", "start": 0.0, "end": 2.0},
                {"text": "Second video segment", "start": 2.0, "end": 4.0},
            ],
            ["images/test-video-frame.jpg"],
        )

    monkeypatch.setattr(process_module, "_process_video", fake_process_video)

    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-video-{uuid.uuid4().hex[:8]}",
                content_type="video",
                source_type="upload",
                file_path="uploads/video-test.mp4",
            )
            db.add(content)
            await db.commit()
            content_id = str(content.id)

        process_resp = await client.post(f"/api/contents/{content_id}/process")
        assert process_resp.status_code == 200
        assert process_resp.json()["processing_status"] in {"completed", "chunked"}

        chunks_resp = await client.get(f"/api/contents/{content_id}/chunks")
        assert chunks_resp.status_code == 200
        chunks = chunks_resp.json()["chunks"]
        assert [chunk["chunk_type"] for chunk in chunks] == ["text", "text", "image"]
        assert chunks[0]["chunk_text"] == "First video segment"
        assert chunks[1]["time_start"] == 2.0
        assert chunks[2]["image_path"] == "images/test-video-frame.jpg"

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == uuid.UUID(content_id)))
            refreshed = result.scalar_one()
            assert refreshed.text_content == "First video segment Second video segment"
    finally:
        settings.file_storage_root = old_storage_root


@pytest.mark.asyncio
async def test_docx_processing_preserves_structure_metadata(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    from docx import Document

    settings = get_settings()
    doc_dir = tmp_path / "uploads"
    doc_dir.mkdir()
    doc_path = doc_dir / "structured.docx"

    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Moyuan"
    table.cell(1, 1).text = "42"
    doc.save(doc_path)

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(tmp_path)
    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-docx-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="upload",
                file_path="uploads/structured.docx",
            )
            db.add(content)
            await db.commit()
            content_id = str(content.id)

        process_resp = await client.post(f"/api/contents/{content_id}/process")
        assert process_resp.status_code == 200

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == uuid.UUID(content_id)))
            refreshed = result.scalar_one()
            structure = refreshed.extra_meta["document_structure"]
            assert structure["format"] == "docx"
            assert structure["headings"][0]["text"] == "Chapter One"
            assert structure["headings"][0]["level"] == 1
            assert structure["tables"][0]["row_count"] == 2
            assert "Body paragraph" in refreshed.text_content
    finally:
        settings.file_storage_root = old_storage_root


@pytest.mark.asyncio
async def test_xlsx_processing_preserves_sheet_metadata(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Content
    from openpyxl import Workbook

    settings = get_settings()
    doc_dir = tmp_path / "uploads"
    doc_dir.mkdir()
    xlsx_path = doc_dir / "structured.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Facts"
    ws.append(["Name", "Value"])
    ws.append(["Moyuan", 42])
    wb.save(xlsx_path)

    old_storage_root = settings.file_storage_root
    settings.file_storage_root = str(tmp_path)
    try:
        async with async_session_factory() as db:
            content = Content(
                title=f"pytest-xlsx-{uuid.uuid4().hex[:8]}",
                content_type="doc",
                source_type="upload",
                file_path="uploads/structured.xlsx",
            )
            db.add(content)
            await db.commit()
            content_id = str(content.id)

        process_resp = await client.post(f"/api/contents/{content_id}/process")
        assert process_resp.status_code == 200

        async with async_session_factory() as db:
            result = await db.execute(select(Content).where(Content.id == uuid.UUID(content_id)))
            refreshed = result.scalar_one()
            structure = refreshed.extra_meta["document_structure"]
            assert structure["format"] == "xlsx"
            assert structure["sheets"][0]["name"] == "Facts"
            assert structure["sheets"][0]["non_empty_rows"] == 2
            assert structure["sheets"][0]["preview"][1] == ["Moyuan", "42"]
            assert "Moyuan | 42" in refreshed.text_content
    finally:
        settings.file_storage_root = old_storage_root


@pytest.mark.asyncio
async def test_embedding_stats(client: AsyncClient):
    resp = await client.get("/api/embeddings/stats")
    assert resp.status_code == 200
    data = resp.json()
    assert "total_text_contents" in data


@pytest.mark.asyncio
async def test_provider_diagnostics_reports_runtime_and_binding_status(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import FunctionBindingConfig, ProviderConfig
    from sqlalchemy import delete

    provider_name = f"pytest-provider-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        await db.execute(delete(FunctionBindingConfig).where(FunctionBindingConfig.function == "qa"))
        provider = ProviderConfig(
            name=provider_name,
            provider_type="openai",
            base_url="https://example.com/v1",
            is_active=True,
        )
        db.add(provider)
        await db.flush()
        db.add(FunctionBindingConfig(function="qa", provider_id=provider.id, model="gpt-test"))
        await db.commit()

    resp = await client.get("/api/providers/diagnostics")
    assert resp.status_code == 200
    data = resp.json()

    assert {check["key"] for check in data["checks"]} >= {"trafilatura", "playwright", "ffmpeg", "faster_whisper"}

    qa_binding = next(item for item in data["bindings"] if item["function"] == "qa")
    assert qa_binding["ok"] is True
    assert qa_binding["provider_name"] == provider_name
    assert qa_binding["model"] == "gpt-test"

    ocr_binding = next(item for item in data["bindings"] if item["function"] == "ocr")
    assert "ok" in ocr_binding
    assert "detail" in ocr_binding


@pytest.mark.asyncio
async def test_function_bindings_validate_provider_and_clean_model(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import FunctionBindingConfig, ProviderConfig
    from sqlalchemy import delete, select

    missing_provider_resp = await client.put(
        "/api/providers/bindings",
        json={
            "bindings": {
                "qa": {
                    "function": "qa",
                    "provider_id": str(uuid.uuid4()),
                    "model": "gpt-test",
                }
            }
        },
    )
    assert missing_provider_resp.status_code == 404

    async with async_session_factory() as db:
        await db.execute(delete(FunctionBindingConfig).where(FunctionBindingConfig.function == "qa"))
        provider = ProviderConfig(
            name=f"pytest-binding-provider-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            base_url="https://example.com/v1",
            is_active=True,
        )
        db.add(provider)
        await db.commit()
        provider_id = str(provider.id)

    trimmed_resp = await client.put(
        "/api/providers/bindings",
        json={
            "bindings": {
                "qa": {
                    "function": "qa",
                    "provider_id": provider_id,
                    "model": "  gpt-trimmed  ",
                }
            }
        },
    )
    assert trimmed_resp.status_code == 200
    qa_binding = trimmed_resp.json()["bindings"]["qa"]
    assert qa_binding["provider_id"] == provider_id
    assert qa_binding["model"] == "gpt-trimmed"

    empty_model_resp = await client.put(
        "/api/providers/bindings",
        json={
            "bindings": {
                "qa": {
                    "function": "qa",
                    "provider_id": provider_id,
                    "model": "  ",
                }
            }
        },
    )
    assert empty_model_resp.status_code == 200
    assert empty_model_resp.json()["bindings"]["qa"]["model"] is None

    async with async_session_factory() as db:
        result = await db.execute(select(FunctionBindingConfig).where(FunctionBindingConfig.function == "qa"))
        stored = result.scalar_one()
        assert str(stored.provider_id) == provider_id
        assert stored.model is None


@pytest.mark.asyncio
async def test_provider_api_key_is_masked_until_revealed(client: AsyncClient):
    raw_key = f"sk-test-{uuid.uuid4().hex}"
    resp = await client.post(
        "/api/providers",
        json={
            "name": f"pytest-key-{uuid.uuid4().hex[:8]}",
            "provider_type": "openai",
            "base_url": "  https://example.com/v1  ",
            "api_key": f"  {raw_key}  ",
            "default_models": {"qa": "  gpt-qa  ", "ocr": "  "},
        },
    )
    assert resp.status_code == 201
    provider = resp.json()
    provider_id = provider["id"]
    assert provider["api_key_masked"] != raw_key
    assert raw_key not in str(provider)
    assert provider["base_url"] == "https://example.com/v1"
    assert provider["default_models"] == {"qa": "gpt-qa"}

    list_resp = await client.get("/api/providers")
    assert list_resp.status_code == 200
    assert raw_key not in str(list_resp.json())

    reveal_resp = await client.get(f"/api/providers/{provider_id}/api-key")
    assert reveal_resp.status_code == 200
    assert reveal_resp.json()["api_key"] == raw_key

    clear_resp = await client.put(
        f"/api/providers/{provider_id}",
        json={"api_key": "", "base_url": " ", "default_models": {"qa": " "}},
    )
    assert clear_resp.status_code == 200
    assert clear_resp.json()["api_key_masked"] is None
    assert clear_resp.json()["base_url"] is None
    assert clear_resp.json()["default_models"] is None

    cleared_key_resp = await client.get(f"/api/providers/{provider_id}/api-key")
    assert cleared_key_resp.status_code == 200
    assert cleared_key_resp.json()["api_key"] is None


@pytest.mark.asyncio
async def test_delete_provider_clears_bindings_and_brain_config(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, FunctionBindingConfig, ProviderConfig
    from sqlalchemy import delete, select

    async with async_session_factory() as db:
        await db.execute(delete(FunctionBindingConfig).where(FunctionBindingConfig.function == "qa"))
        provider = ProviderConfig(
            name=f"pytest-delete-provider-{uuid.uuid4().hex[:8]}",
            provider_type="openai",
            base_url="https://example.com/v1",
            is_active=True,
        )
        brain = Brain(
            name=f"pytest-delete-provider-brain-{uuid.uuid4().hex[:8]}",
            config={"provider_id": "", "qa_model": "gpt-before-delete"},
        )
        db.add_all([provider, brain])
        await db.flush()
        brain.config = {"provider_id": str(provider.id), "qa_model": "gpt-before-delete"}
        db.add(FunctionBindingConfig(function="qa", provider_id=provider.id, model="gpt-before-delete"))
        await db.commit()
        provider_id = str(provider.id)
        brain_id = brain.id

    resp = await client.delete(f"/api/providers/{provider_id}")
    assert resp.status_code == 204

    async with async_session_factory() as db:
        binding_result = await db.execute(select(FunctionBindingConfig).where(FunctionBindingConfig.function == "qa"))
        binding = binding_result.scalar_one()
        refreshed_brain = await db.get(Brain, brain_id)
        assert binding.provider_id is None
        assert binding.model == "gpt-before-delete"
        assert refreshed_brain.config == {"qa_model": "gpt-before-delete"}


@pytest.mark.asyncio
async def test_relation_suggestions_return_top_unrelated_similar_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentRelation

    def vector(first: float, second: float = 0.0) -> list[float]:
        values = [0.0] * 4096
        values[3072] = first
        values[3073] = second
        return values

    brain_id = uuid.uuid4()
    async with async_session_factory() as db:
        source = Content(
            title=f"pytest-source-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="source",
            embedding=vector(1.0),
            processing_status="completed",
            brain_id=brain_id,
        )
        close = Content(
            title=f"pytest-close-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="close",
            embedding=vector(0.99, 0.01),
            processing_status="completed",
            brain_id=brain_id,
        )
        already_related = Content(
            title=f"pytest-related-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="related",
            embedding=vector(0.98, 0.02),
            processing_status="completed",
            brain_id=brain_id,
        )
        no_embedding = Content(
            title=f"pytest-no-embedding-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="plain",
            processing_status="completed",
            brain_id=brain_id,
        )
        db.add_all([source, close, already_related, no_embedding])
        await db.flush()
        db.add(
            ContentRelation(
                source_id=source.id,
                target_id=already_related.id,
                relation_type="similar",
            )
        )
        await db.commit()
        source_id = str(source.id)
        close_id = str(close.id)
        related_id = str(already_related.id)
        no_embedding_id = str(no_embedding.id)

    resp = await client.get(f"/api/relations/suggestions?content_id={source_id}&limit=5")
    assert resp.status_code == 200
    data = resp.json()
    ids = [item["id"] for item in data]
    assert close_id in ids
    assert related_id not in ids
    assert all(item["similarity"] <= 1 for item in data)

    empty_resp = await client.get(f"/api/relations/suggestions?content_id={no_embedding_id}")
    assert empty_resp.status_code == 200
    assert empty_resp.json() == []


@pytest.mark.asyncio
async def test_relation_create_rejects_cross_brain_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content

    async with async_session_factory() as db:
        source = Content(
            title=f"pytest-relation-brain-a-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="source",
            brain_id=uuid.uuid4(),
        )
        target = Content(
            title=f"pytest-relation-brain-b-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="target",
            brain_id=uuid.uuid4(),
        )
        db.add_all([source, target])
        await db.commit()
        source_id = str(source.id)
        target_id = str(target.id)

    resp = await client.post(
        "/api/relations",
        json={"source_id": source_id, "target_id": target_id, "relation_type": "reference"},
    )
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_relations_hide_deleted_content(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentRelation

    source_title = f"relation-visible-{uuid.uuid4().hex[:8]}"
    async with async_session_factory() as db:
        source = Content(
            title=source_title,
            content_type="doc",
            source_type="manual",
            text_content="source",
            processing_status="completed",
        )
        deleted_target = Content(
            title=f"relation-deleted-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="deleted",
            processing_status="completed",
            is_deleted=True,
        )
        deleted_source = Content(
            title=f"relation-deleted-source-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="deleted source",
            processing_status="completed",
            is_deleted=True,
        )
        db.add_all([source, deleted_target, deleted_source])
        await db.flush()
        db.add_all([
            ContentRelation(source_id=source.id, target_id=deleted_target.id, relation_type="reference"),
            ContentRelation(source_id=source.id, target_id=deleted_target.id, relation_type="series"),
        ])
        await db.commit()
        source_id = str(source.id)
        deleted_source_id = str(deleted_source.id)

    list_resp = await client.get(f"/api/relations?content_id={source_id}")
    series_resp = await client.get(f"/api/relations/series?content_id={source_id}")
    deleted_source_resp = await client.get(f"/api/relations?content_id={deleted_source_id}")

    assert list_resp.status_code == 200
    assert list_resp.json() == []
    assert series_resp.status_code == 200
    assert series_resp.json()["items"] == [{"id": source_id, "title": source_title, "sort_order": 0}]
    assert series_resp.json()["total"] == 1
    assert deleted_source_resp.status_code == 404


@pytest.mark.asyncio
async def test_related_content_includes_explicit_graph_relations_without_embeddings(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentRelation

    async with async_session_factory() as db:
        source = Content(
            title=f"pytest-related-source-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="source without embedding",
            processing_status="completed",
        )
        reference = Content(
            title=f"pytest-reference-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="reference without embedding",
            processing_status="completed",
        )
        series = Content(
            title=f"pytest-series-{uuid.uuid4().hex[:8]}",
            content_type="video",
            source_type="manual",
            text_content="series without embedding",
            processing_status="completed",
        )
        db.add_all([source, reference, series])
        await db.flush()
        db.add_all([
            ContentRelation(source_id=source.id, target_id=reference.id, relation_type="reference"),
            ContentRelation(source_id=source.id, target_id=series.id, relation_type="series"),
        ])
        await db.commit()
        source_id = str(source.id)
        reference_id = str(reference.id)
        series_id = str(series.id)

    resp = await client.get(f"/api/ai/related/{source_id}?top_k=5")
    assert resp.status_code == 200
    related = resp.json()["related"]
    ids = [item["id"] for item in related]
    assert ids[:2] == [series_id, reference_id]
    assert related[0]["relation_type"] == "series"
    assert related[1]["relation_type"] == "reference"
    assert related[0]["relation_bonus"] > related[1]["relation_bonus"]


# ── Analytics ──

@pytest.mark.asyncio
async def test_analytics_overview(client: AsyncClient):
    resp = await client.get("/api/analytics/overview")
    assert resp.status_code == 200
    data = resp.json()
    assert "total_contents" in data


@pytest.mark.asyncio
async def test_search_logs_and_analytics_are_brain_scoped(client: AsyncClient):
    from datetime import datetime, timezone
    from app.core.database import async_session_factory
    from app.models.models import Brain, Content, ContentChunk, SearchLog

    brain_a = uuid.uuid4()
    brain_b = uuid.uuid4()
    query_a = f"brain-query-a-{uuid.uuid4().hex[:8]}"
    query_b = f"brain-query-b-{uuid.uuid4().hex[:8]}"

    async with async_session_factory() as db:
        db.add_all([
            Brain(id=brain_a, name=f"search-brain-a-{uuid.uuid4().hex[:8]}"),
            Brain(id=brain_b, name=f"search-brain-b-{uuid.uuid4().hex[:8]}"),
        ])
        content_a = Content(
            title=f"pytest-search-brain-a-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content=query_a,
            processing_status="completed",
            brain_id=brain_a,
        )
        db.add(content_a)
        await db.flush()
        db.add(ContentChunk(content_id=content_a.id, chunk_index=0, chunk_type="text", chunk_text=query_a))
        db.add_all([
            SearchLog(query=query_a, result_count=1, brain_id=brain_a, created_at=datetime.now(timezone.utc)),
            SearchLog(query=query_b, result_count=1, brain_id=brain_b, created_at=datetime.now(timezone.utc)),
        ])
        await db.commit()

    search_resp = await client.post(
        "/api/search",
        json={"query": query_a, "top_k": 5, "brain_id": str(brain_a), "enable_vector": False},
    )
    assert search_resp.status_code == 200
    assert search_resp.json()["total"] >= 1

    history_resp = await client.get(f"/api/search/history?brain_id={brain_a}&page_size=100")
    assert history_resp.status_code == 200
    history = history_resp.json()["items"]
    assert any(item["query"] == query_a and item["brain_id"] == str(brain_a) for item in history)
    assert all(item["brain_id"] in {str(brain_a), None} for item in history)

    trends_resp = await client.get(f"/api/analytics/search-trends?brain_id={brain_a}&limit=10")
    assert trends_resp.status_code == 200
    payload = trends_resp.json()
    assert any(item["query"] == query_a for item in payload["trends"])
    assert all(item["query"] != query_b for item in payload["trends"])
    assert "daily" in payload


@pytest.mark.asyncio
async def test_embedding_reindex_is_brain_scoped(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, Content, ContentChunk

    brain_a = uuid.uuid4()
    brain_b = uuid.uuid4()

    async with async_session_factory() as db:
        db.add_all([
            Brain(id=brain_a, name=f"embed-brain-a-{uuid.uuid4().hex[:8]}"),
            Brain(id=brain_b, name=f"embed-brain-b-{uuid.uuid4().hex[:8]}"),
        ])
        content_a = Content(
            title=f"pytest-embed-a-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="a",
            embedding=[0.1] * 4096,
            embedding_type="text",
            brain_id=brain_a,
        )
        content_b = Content(
            title=f"pytest-embed-b-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="b",
            embedding=[0.2] * 4096,
            embedding_type="text",
            brain_id=brain_b,
        )
        db.add_all([content_a, content_b])
        await db.flush()
        db.add_all([
            ContentChunk(content_id=content_a.id, chunk_index=0, chunk_type="text", chunk_text="a"),
            ContentChunk(content_id=content_b.id, chunk_index=0, chunk_type="text", chunk_text="b"),
        ])
        await db.commit()
        content_a_id = content_a.id
        content_b_id = content_b.id

    resp = await client.post(f"/api/embeddings/reindex?brain_id={brain_a}")
    assert resp.status_code == 200
    assert resp.json()["cleared"] == 1

    async with async_session_factory() as db:
        a_after = await db.get(Content, content_a_id)
        b_after = await db.get(Content, content_b_id)
        chunks_a = await db.execute(select(ContentChunk).where(ContentChunk.content_id == content_a_id))
        chunks_b = await db.execute(select(ContentChunk).where(ContentChunk.content_id == content_b_id))
        assert a_after.embedding is None
        assert b_after.embedding is not None
        assert chunks_a.scalars().all() == []
        assert len(chunks_b.scalars().all()) == 1


@pytest.mark.asyncio
async def test_scoped_search_analytics_and_embedding_reject_unknown_brain(client: AsyncClient):
    missing_brain = uuid.uuid4()

    search_resp = await client.post(
        "/api/search",
        json={"query": "missing brain", "brain_id": str(missing_brain), "enable_vector": False},
    )
    history_resp = await client.get(f"/api/search/history?brain_id={missing_brain}")
    analytics_resp = await client.get(f"/api/analytics/overview?brain_id={missing_brain}")
    embedding_stats_resp = await client.get(f"/api/embeddings/stats?brain_id={missing_brain}")
    embedding_reindex_resp = await client.post(f"/api/embeddings/reindex?brain_id={missing_brain}")

    assert search_resp.status_code == 404
    assert history_resp.status_code == 404
    assert analytics_resp.status_code == 404
    assert embedding_stats_resp.status_code == 404
    assert embedding_reindex_resp.status_code == 404


# ── Backup ──

@pytest.mark.asyncio
async def test_backup_list(client: AsyncClient):
    resp = await client.get("/api/backup")
    assert resp.status_code == 200
    assert "backups" in resp.json()


@pytest.mark.asyncio
async def test_backup_export_writes_manifest_and_redacts_provider_api_keys(client: AsyncClient, tmp_path, monkeypatch):
    from app.core.config import get_settings
    from app.core.crypto import crypto_service
    from app.core.database import async_session_factory
    from app.models.models import Brain, ProviderConfig
    from pathlib import Path
    import json
    import zipfile
    import app.api.backup as backup_module

    settings = get_settings()
    old_storage_root = settings.file_storage_root
    old_backup_storage_root = backup_module.settings.file_storage_root
    storage_root = tmp_path / "storage"
    storage_root.mkdir()
    settings.file_storage_root = str(storage_root)
    backup_module.settings.file_storage_root = str(storage_root)

    raw_key = f"sk-backup-{uuid.uuid4().hex}"
    encrypted_key = crypto_service.encrypt(raw_key)
    backup_brain_name = f"pytest-backup-brain-{uuid.uuid4().hex[:8]}"

    def fake_run(cmd, env=None, check=False, capture_output=True, timeout=60):
        sql_path = cmd[cmd.index("-f") + 1]
        Path(sql_path).write_text(
            "\n".join([
                "COPY public.provider_configs (id, name, provider_type, base_url, api_key_encrypted, default_models, extra_params, is_active, created_at, updated_at) FROM stdin;",
                f"{uuid.uuid4()}\tProvider\topenai\thttps://example.test/v1\t{encrypted_key}\t\\N\t\\N\tt\t\\N\t\\N",
                r"\.",
                "",
            ]),
            encoding="utf-8",
        )
        class Result:
            returncode = 0
            stdout = b""
            stderr = b""
        return Result()

    monkeypatch.setattr(backup_module.subprocess, "run", fake_run)

    try:
        async with async_session_factory() as db:
            db.add(
                ProviderConfig(
                    name=f"pytest-backup-provider-{uuid.uuid4().hex[:8]}",
                    provider_type="openai",
                    base_url="https://example.test/v1",
                    api_key_encrypted=encrypted_key,
                    is_active=True,
                )
            )
            db.add(
                Brain(
                    name=backup_brain_name,
                    description="backup model config",
                    icon="brain",
                    config={
                        "provider_id": str(uuid.uuid4()),
                        "summarize_model": "summary-export-model",
                        "qa_model": "qa-export-model",
                        "quiz_model": "quiz-export-model",
                        "judge_model": "judge-export-model",
                        "embedding_model": "embed-export-model",
                    },
                )
            )
            await db.commit()

        resp = await client.post("/api/backup/export")
        assert resp.status_code == 200
        data = resp.json()
        assert data["api_keys_included"] is False
        assert data["format_version"] == 1

        zip_path = backup_module._get_backup_dir() / data["filename"]
        with zipfile.ZipFile(zip_path, "r") as zf:
            names = set(zf.namelist())
            assert {"manifest.json", "config/providers.json", "config/function_bindings.json", "config/brains.json", "database.sql"} <= names
            manifest = zf.read("manifest.json").decode("utf-8")
            providers_json = zf.read("config/providers.json").decode("utf-8")
            brains_json = zf.read("config/brains.json").decode("utf-8")
            database_sql = zf.read("database.sql").decode("utf-8")

        brains_payload = json.loads(brains_json)
        exported_brain = next(brain for brain in brains_payload if brain["name"] == backup_brain_name)
        assert exported_brain["config"]["summarize_model"] == "summary-export-model"
        assert exported_brain["config"]["qa_model"] == "qa-export-model"
        assert exported_brain["config"]["judge_model"] == "judge-export-model"

        combined = "\n".join([manifest, providers_json, brains_json, database_sql])
        assert raw_key not in combined
        assert encrypted_key not in combined
        assert '"api_keys_included": false' in manifest
        assert '"brain_configs": true' in manifest
        assert '"api_key": null' in providers_json
        assert '"api_key_encrypted": null' in providers_json
        assert "\t\\N\t" in database_sql
    finally:
        settings.file_storage_root = old_storage_root
        backup_module.settings.file_storage_root = old_backup_storage_root


@pytest.mark.asyncio
async def test_backup_restore_restores_files_and_rejects_unsafe_names(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Brain, FunctionBindingConfig, ProviderConfig
    from sqlalchemy import select
    import json
    import zipfile
    import app.api.backup as backup_module

    settings = get_settings()
    old_storage_root = settings.file_storage_root
    old_backup_storage_root = backup_module.settings.file_storage_root
    storage_root = tmp_path / "storage"
    settings.file_storage_root = str(storage_root)
    backup_module.settings.file_storage_root = str(storage_root)
    backup_dir = backup_module._get_backup_dir()
    try:
        zip_path = backup_dir / "restore-test.zip"
        provider_id = uuid.uuid4()
        brain_id = uuid.uuid4()
        binding_function = f"backup_restore_{uuid.uuid4().hex[:8]}"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(
                "manifest.json",
                json.dumps(
                    {
                        "format_version": 1,
                        "contains": {"brain_configs": True},
                        "security": {"api_keys_included": False},
                    }
                ),
            )
            zf.writestr(
                "config/brains.json",
                json.dumps([
                    {
                        "id": str(brain_id),
                        "name": "restore-brain",
                        "description": "restored",
                        "icon": "brain",
                        "config": {"provider_id": str(provider_id), "qa_model": "qa"},
                    }
                ]),
            )
            zf.writestr(
                "config/providers.json",
                json.dumps([
                    {
                        "id": str(provider_id),
                        "name": "restore-provider",
                        "provider_type": "openai",
                        "base_url": "https://restore.example/v1",
                        "default_models": {"qa": "qa"},
                        "extra_params": {"temperature": 0},
                        "is_active": True,
                    }
                ]),
            )
            zf.writestr(
                "config/function_bindings.json",
                json.dumps([
                    {
                        "function": binding_function,
                        "provider_id": str(provider_id),
                        "model": "restore-model",
                        "extra_params": {"top_p": 1},
                    }
                ]),
            )
            zf.writestr("database.sql", "")
            zf.writestr("files/uploads/restored.txt", "restored file")

        inspect_resp = await client.get("/api/backup/restore-test.zip/inspect")
        assert inspect_resp.status_code == 200
        inspect_data = inspect_resp.json()
        assert inspect_data["format_version"] == 1
        assert inspect_data["api_keys_included"] is False
        assert inspect_data["brain_configs"] == 1
        assert inspect_data["file_count"] == 1
        assert inspect_data["has_database_sql"] is False
        assert inspect_data["config_preview"] == {
            "providers": {"new": 1, "overwrite": 0, "invalid": 0},
            "function_bindings": {"new": 1, "overwrite": 0, "invalid": 0},
            "brains": {"new": 1, "overwrite": 0, "invalid": 0},
        }

        resp = await client.post("/api/backup/restore-test.zip/restore")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "restored"
        assert data["restored_files"] == 1
        assert data["database_status"] == "missing"
        assert data["config_status"] == "restored"
        assert data["restored_config"] == {"providers": 1, "function_bindings": 1, "brains": 1}
        assert (storage_root / "uploads" / "restored.txt").read_text(encoding="utf-8") == "restored file"

        async with async_session_factory() as db:
            provider = await db.get(ProviderConfig, provider_id)
            brain = await db.get(Brain, brain_id)
            binding_result = await db.execute(
                select(FunctionBindingConfig).where(FunctionBindingConfig.function == binding_function)
            )
            binding = binding_result.scalar_one()
            assert provider is not None
            assert provider.name == "restore-provider"
            assert provider.api_key_encrypted is None
            assert brain is not None
            assert brain.config["qa_model"] == "qa"
            assert binding.provider_id == provider_id
            assert binding.model == "restore-model"

        unsafe_zip = backup_dir / "unsafe-restore.zip"
        with zipfile.ZipFile(unsafe_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("../evil.txt", "nope")
        unsafe_restore_resp = await client.post("/api/backup/unsafe-restore.zip/restore")
        assert unsafe_restore_resp.status_code == 400

        unsafe_resp = await client.delete("/api/backup/restore-test.txt")
        assert unsafe_resp.status_code == 400
    finally:
        settings.file_storage_root = old_storage_root
        backup_module.settings.file_storage_root = old_backup_storage_root


@pytest.mark.asyncio
async def test_backup_restore_modes_can_restore_config_or_files_only(client: AsyncClient, tmp_path):
    from app.core.config import get_settings
    from app.core.database import async_session_factory
    from app.models.models import Brain, FunctionBindingConfig, ProviderConfig
    import json
    import zipfile
    import app.api.backup as backup_module

    settings = get_settings()
    old_storage_root = settings.file_storage_root
    old_backup_storage_root = backup_module.settings.file_storage_root
    storage_root = tmp_path / "storage"
    settings.file_storage_root = str(storage_root)
    backup_module.settings.file_storage_root = str(storage_root)
    backup_dir = backup_module._get_backup_dir()
    provider_id = uuid.uuid4()
    brain_id = uuid.uuid4()
    binding_function = f"restore_mode_binding_{uuid.uuid4().hex[:8]}"

    try:
        async with async_session_factory() as db:
            db.add_all([
                ProviderConfig(
                    id=provider_id,
                    name="old-mode-provider",
                    provider_type="openai",
                    base_url="https://old.example/v1",
                    is_active=True,
                ),
                Brain(id=brain_id, name="old-mode-brain", config={"summarize_model": "old-summary"}),
                FunctionBindingConfig(function=binding_function, model="old-model"),
            ])
            await db.commit()

        zip_path = backup_dir / "restore-mode-test.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("manifest.json", json.dumps({"format_version": 1, "security": {"api_keys_included": False}}))
            zf.writestr(
                "config/providers.json",
                json.dumps([
                    {
                        "id": str(provider_id),
                        "name": "mode-provider",
                        "provider_type": "openai",
                        "base_url": "https://mode.example/v1",
                        "is_active": True,
                    },
                    {"id": "not-a-uuid", "name": "bad-provider", "provider_type": "openai"},
                ]),
            )
            zf.writestr(
                "config/brains.json",
                json.dumps([
                    {
                        "id": str(brain_id),
                        "name": "mode-brain",
                        "config": {"provider_id": str(provider_id), "summarize_model": "mode-summary"},
                    },
                    {"id": "not-a-uuid", "name": "bad-brain"},
                ]),
            )
            zf.writestr(
                "config/function_bindings.json",
                json.dumps([
                    {"function": binding_function, "provider_id": str(provider_id), "model": "mode-model"},
                    {"function": ""},
                ]),
            )
            zf.writestr("database.sql", "select 1;")
            zf.writestr("files/uploads/mode-restored.txt", "mode file")

        inspect_resp = await client.get("/api/backup/restore-mode-test.zip/inspect")
        assert inspect_resp.status_code == 200
        assert inspect_resp.json()["config_preview"] == {
            "providers": {"new": 0, "overwrite": 1, "invalid": 1},
            "function_bindings": {"new": 0, "overwrite": 1, "invalid": 1},
            "brains": {"new": 0, "overwrite": 1, "invalid": 1},
        }

        config_resp = await client.post("/api/backup/restore-mode-test.zip/restore?mode=config")
        assert config_resp.status_code == 200
        config_data = config_resp.json()
        assert config_data["mode"] == "config"
        assert config_data["restored_files"] == 0
        assert config_data["database_status"] == "skipped"
        assert config_data["config_status"] == "restored"
        assert config_data["restored_config"] == {"providers": 1, "function_bindings": 1, "brains": 1}
        assert not (storage_root / "uploads" / "mode-restored.txt").exists()

        async with async_session_factory() as db:
            provider = await db.get(ProviderConfig, provider_id)
            brain = await db.get(Brain, brain_id)
            assert provider is not None
            assert provider.name == "mode-provider"
            assert brain is not None
            assert brain.config["summarize_model"] == "mode-summary"
            binding_result = await db.execute(
                select(FunctionBindingConfig).where(FunctionBindingConfig.function == binding_function)
            )
            assert binding_result.scalar_one().model == "mode-model"

        files_resp = await client.post("/api/backup/restore-mode-test.zip/restore?mode=files")
        assert files_resp.status_code == 200
        files_data = files_resp.json()
        assert files_data["mode"] == "files"
        assert files_data["restored_files"] == 1
        assert files_data["database_status"] == "skipped"
        assert files_data["config_status"] == "skipped"
        assert (storage_root / "uploads" / "mode-restored.txt").read_text(encoding="utf-8") == "mode file"
    finally:
        settings.file_storage_root = old_storage_root
        backup_module.settings.file_storage_root = old_backup_storage_root


# ── 异常场景 ──

@pytest.mark.asyncio
async def test_get_nonexistent_file(client: AsyncClient):
    resp = await client.get("/api/files/00000000-0000-0000-0000-000000000000")
    assert resp.status_code == 404


@pytest.mark.asyncio
async def test_empty_search(client: AsyncClient):
    resp = await client.post("/api/search", json={"query": "", "top_k": 5})
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_batch_content_status_returns_chunk_stats_and_validates_scope(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Content, ContentChunk

    brain_resp = await client.post("/api/brains", json={"name": f"status-batch-{uuid.uuid4().hex[:8]}"})
    assert brain_resp.status_code == 200
    brain_id = uuid.UUID(brain_resp.json()["id"])

    async with async_session_factory() as db:
        content = Content(
            title=f"status-batch-content-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="manual",
            text_content="chunked content",
            processing_status="completed",
            brain_id=brain_id,
        )
        db.add(content)
        await db.flush()
        db.add_all([
            ContentChunk(content_id=content.id, chunk_index=0, chunk_type="text", chunk_text="a", embedding=[0.0] * 4096),
            ContentChunk(content_id=content.id, chunk_index=1, chunk_type="image", image_path="images/a.png"),
        ])
        await db.commit()
        content_id = str(content.id)

    resp = await client.post(
        "/api/contents/status-batch",
        json={"ids": [content_id], "brain_id": str(brain_id)},
    )
    bad_id_resp = await client.post("/api/contents/status-batch", json={"ids": ["not-a-uuid"]})
    missing_brain_resp = await client.post(
        "/api/contents/status-batch",
        json={"ids": [], "brain_id": str(uuid.uuid4())},
    )

    assert resp.status_code == 200
    status = resp.json()["items"][content_id]
    assert status["processing_status"] == "completed"
    assert status["has_text"] is True
    assert status["chunk_count"] == 2
    assert status["text_chunks"] == 1
    assert status["image_chunks"] == 1
    assert status["embedded_chunks"] == 1
    assert bad_id_resp.status_code == 400
    assert missing_brain_resp.status_code == 404


@pytest.mark.asyncio
async def test_processing_center_summarizes_scope_and_latest_tasks(client: AsyncClient):
    from app.core.database import async_session_factory
    from app.models.models import Brain, Content, ContentChunk, ProcessingTask

    async with async_session_factory() as db:
        brain = Brain(name=f"processing-center-{uuid.uuid4().hex[:8]}")
        other_brain = Brain(name=f"processing-center-other-{uuid.uuid4().hex[:8]}")
        db.add_all([brain, other_brain])
        await db.flush()

        active = Content(
            title=f"processing-active-{uuid.uuid4().hex[:8]}",
            content_type="pdf",
            source_type="upload",
            processing_status="embedding",
            brain_id=brain.id,
        )
        failed = Content(
            title=f"processing-failed-{uuid.uuid4().hex[:8]}",
            content_type="doc",
            source_type="upload",
            processing_status="failed",
            processing_error="pytest failure",
            brain_id=brain.id,
        )
        completed = Content(
            title=f"processing-completed-{uuid.uuid4().hex[:8]}",
            content_type="note",
            source_type="manual",
            processing_status="completed",
            brain_id=other_brain.id,
        )
        db.add_all([active, failed, completed])
        await db.flush()
        db.add_all([
            ContentChunk(content_id=active.id, chunk_index=0, chunk_type="text", chunk_text="a", embedding=[0.0] * 4096),
            ContentChunk(content_id=active.id, chunk_index=1, chunk_type="text", chunk_text="b"),
            ProcessingTask(content_id=active.id, task_type="embed", status="processing", progress=35),
            ProcessingTask(content_id=failed.id, task_type="parse", status="failed", progress=0, error_message="task failed"),
        ])
        await db.commit()
        brain_id = str(brain.id)
        active_id = str(active.id)
        failed_id = str(failed.id)

    active_resp = await client.get(f"/api/contents/processing-center?brain_id={brain_id}&group=active")
    failed_resp = await client.get(f"/api/contents/processing-center?brain_id={brain_id}&group=failed")
    missing_resp = await client.get(f"/api/contents/processing-center?brain_id={uuid.uuid4()}")

    assert active_resp.status_code == 200
    active_data = active_resp.json()
    assert active_data["summary"]["total"] == 2
    assert active_data["summary"]["active"] == 1
    assert active_data["summary"]["failed"] == 1
    assert active_data["tasks"]["processing"] == 1
    assert active_data["items"][0]["id"] == active_id
    assert active_data["items"][0]["chunk_count"] == 2
    assert active_data["items"][0]["embedded_chunks"] == 1
    assert active_data["items"][0]["latest_task"]["task_type"] == "embed"

    assert failed_resp.status_code == 200
    failed_items = failed_resp.json()["items"]
    assert len(failed_items) == 1
    assert failed_items[0]["id"] == failed_id
    assert failed_items[0]["latest_task"]["error_message"] == "task failed"
    assert missing_resp.status_code == 404


@pytest.mark.asyncio
async def test_invalid_content_id(client: AsyncClient):
    resp = await client.get("/api/contents/not-a-uuid/status")
    assert resp.status_code >= 400
